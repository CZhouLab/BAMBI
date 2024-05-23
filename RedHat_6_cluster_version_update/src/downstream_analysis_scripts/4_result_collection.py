
import os
import pandas as pd
from io import StringIO
import time
import re
import argparse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2image import convert_from_path
import subprocess
import os.path
import shutil
import sys
import warnings



def Jobstatus_Check(directory, biomarker_target_gene_type="protein_coding", dataset_name="Customized_Dataset", resubmit=False, HPC_parallel = False):
    ############################################################
    ### parameter need to set by user:

    # biomarker_target_gene_type = "protein_coding"  ## "protein_coding" or "lincRNA"
    #
    # dataset_name = "GSE54456_PC"

    ############################################################


    folder_path = directory

    outlier_remove = True

    folder_list = [folder_path]
    if outlier_remove:
        dataset_name_list = [dataset_name + "_outlier_remove"]
    else:
        dataset_name_list = [dataset_name]

    if not HPC_parallel:
        src_path = "/usr/src/app/src"
    else:
        src_path = directory + "/src"

    gene_type_list = [biomarker_target_gene_type] # ["microarray", "protein_coding", "lincRNA", "CircRNA"]

    DE_type = "wilcox"  ## "wilcox", "DEseq2", "limma"

    High_Corr_Remove_list = [1.0]  # [0.8, 0.9, 1.0]

    parent_CV_partition = 10

    CV_partition = 10

    model_list = ["svm", "knn", "logReg", "bayes"]  # ["svm", "knn", "logReg", "decisionTree", "randomForest", "bayes", "xgboost", "stacking"]

    #######==========================


    # annotation_file = "LncBook_Version2.0_all"  # ["LncBook_Version2.0_all", "gencode_v22", "gencode_v29", "gencode_v37", any path to gtf]

    Target_Metric = "Balanced_Accuracy"


    summary_type = "fresh"  #"fresh", "general"

    intersection_check = str(0) ## 0: not do, 1: do

    Target_range = 0.95


    report_path_record_list = []

    current_partition = 0
    current_parent_partition_folder_path = folder_list[current_partition]



    output_directory = current_parent_partition_folder_path + "/output"
    ML_path = output_directory + "/ML"

    for High_Corr_Remove_threshold in High_Corr_Remove_list:

        if Target_range >= 0 and Target_range <= 1:
            tem_num = str(int(100 * Target_range))
        elif Target_range > 1 and Target_range <= 100:
            tem_num = str(int(Target_range))

        current_time = time.strftime('%Y%m%d%H%M', time.localtime(time.time()))
        if summary_type == "fresh":
            Current_Report_folder_path = ML_path + "/Report_Fresh_" + tem_num + "_" + str(High_Corr_Remove_threshold)
        elif summary_type == "general":
            Current_Report_folder_path = ML_path + "/Report_General_" + tem_num + "_" + str(High_Corr_Remove_threshold)
        # if not os.path.isdir(Current_Report_path):
        #     os.mkdir(Current_Report_path)

        if os.path.isdir(Current_Report_folder_path):
            shutil.rmtree(Current_Report_folder_path)
        os.mkdir(Current_Report_folder_path)


        # if annotation_file == "gencode_v22":  ## ["gencode_v22", "gencode_v29", "gencode_v37", any path to gtf]
        #     gene_name_trans_file = "/nl/umw_chan_zhou/Billy/annotation_file/gencode_v22/gencode_v22_ID_map.txt"
        # elif annotation_file == "gencode_v29":
        #     gene_name_trans_file = "/nl/umw_chan_zhou/Billy/annotation_file/gencode_v29/gencode_v29_ID_map.txt"
        # elif annotation_file == "gencode_v37":
        #     gene_name_trans_file = "/nl/umw_chan_zhou/Billy/annotation_file/gencode_v37/gencode_v37_ID_map.txt"
        # elif annotation_file == "LncBook_Version2.0_all":
        #     gene_name_trans_file = "/nl/umw_chan_zhou/Billy/annotation_file/LncBook_Version2.0_all/LncBook_Version2.0_all_ID_map.txt"
        # else:
        #     create_gene_name_trans_file = subprocess.Popen(["/nl/umw_chan_zhou/Billy/annotation_file/transcript_id2gene_id.gname.gtf.pl", annotation_file],stdout=subprocess.PIPE, )
        #     stdout_value = create_gene_name_trans_file.communicate()[0]
        #     tem_folder_path = src_path + "/annotation_file"
        #     if not os.path.isdir(tem_folder_path):
        #         os.mkdir(tem_folder_path)
        #     create_gene_name_trans_file_output = tem_folder_path + "/customized_ID_map.txt"
        #     logfile = open(create_gene_name_trans_file_output, 'w')
        #     logfile.write(stdout_value.decode('utf-8'))
        #     logfile.close()
        #     gene_name_trans_file = create_gene_name_trans_file_output


        # def add_information_heatmap_all(FPKM_path):
        #
        #     FPKM_df = pd.read_csv(FPKM_path, sep="\t", index_col=0)
        #     FPKM_df_transposition = FPKM_df.T
        #
        #     Label_list = []
        #
        #     for i in range(len(FPKM_df_transposition)):
        #         if str(FPKM_df_transposition.index.values[i]).endswith('_C'):
        #             Label_list.append(0)
        #         elif str(FPKM_df_transposition.index.values[i]).endswith('_T'):
        #             Label_list.append(1)
        #         else:
        #             print("Label not match")
        #
        #     FPKM_df_transposition['label'] = Label_list
        #     FPKM_df_transposition_0 = FPKM_df_transposition[FPKM_df_transposition['label'].isin([0])]
        #     FPKM_df_transposition_1 = FPKM_df_transposition[FPKM_df_transposition['label'].isin([1])]
        #
        #     C_number = len(FPKM_df_transposition_0)
        #
        #     FPKM_df_transposition_0_transform = FPKM_df_transposition_0.drop(['label'], axis=1).T
        #     FPKM_df_transposition_1_transform = FPKM_df_transposition_1.drop(['label'], axis=1).T
        #
        #     FPKM_df = pd.merge(FPKM_df_transposition_0_transform, FPKM_df_transposition_1_transform, left_index=True,
        #                        right_index=True)
        #
        #     return FPKM_df, C_number
        #
        bjobs_str = os.popen('bjobs').read()
        if bjobs_str != '':
            bjobs_trans = StringIO(bjobs_str)
            df_bjobs = pd.read_csv(bjobs_trans, index_col=None)
            JobID_list = df_bjobs['JOBID   USER    STAT  QUEUE      FROM_HOST   EXEC_HOST   JOB_NAME   SUBMIT_TIME'].map(lambda x:x.split()[0]).values.tolist()
            JobStatus_list = df_bjobs['JOBID   USER    STAT  QUEUE      FROM_HOST   EXEC_HOST   JOB_NAME   SUBMIT_TIME'].map(lambda x:x.split()[2]).values.tolist()
            # df_run_status = pd.DataFrame({"JobID":JobID_list, "Status" : JobStatus_list})

        record_df = pd.read_csv(ML_path + '/record.csv', index_col=None)

        ori_record_df = pd.read_csv(ML_path + '/record.csv', index_col=None)

        # os.chdir(DE_output_folder_path)

        # if Heatmap_generation:
        #     Heatmap_pic_folder = Current_Report_folder_path + "/Heatmap_pic"
        #
        #     if os.path.isdir(Heatmap_pic_folder):
        #         shutil.rmtree(Heatmap_pic_folder)
        #     os.mkdir(Heatmap_pic_folder)
        #
        #     output_report_doc = Document()
        #     output_report_doc.add_heading(' HeatMap Overview ', 0)
        #
        #     pic_no = 0
        #
        #
        #
        #     if "protein_coding" in gene_type_list:
        #         tem_pdf_path = DE_output_folder_path + "/PC_heatmap.pdf"
        #         if os.path.isfile(tem_pdf_path):
        #             tem_str = "rm " + tem_pdf_path
        #             os.system(tem_str)
        #
        #         if PC_padj_path != False and PC_FPKM_path != False:
        #             PC_FPKM_df, C_number = add_information_heatmap_all(PC_FPKM_path)
        #
        #             PC_biomarker_FPKM_df_path = DE_output_folder_path + "/PC_biomarker_FPKM_df.txt"
        #             if os.path.isfile(PC_biomarker_FPKM_df_path):
        #                 tem_str = "rm " + PC_biomarker_FPKM_df_path
        #                 os.system(tem_str)
        #
        #             # os.system(tem_str)
        #             PC_FPKM_df.to_csv("./PC_biomarker_FPKM_df.txt", sep="\t", na_rep='NA')
        #             R_script_path = directory + "/src/Heatmap_from_DEseq2_ProteinCodingGene_all.R"
        #             PC_heatmap_all = subprocess.Popen(['Rscript', R_script_path, str(C_number), PC_padj_path,
        #                               "PC_biomarker_FPKM_df.txt", "PC_heatmap.pdf", pvalue_type, pvalue_threshold, PC_min_FPKM_expression])
        #             PC_heatmap_all.communicate()
        #             if os.path.isfile(tem_pdf_path):
        #                 paragraph = output_report_doc.add_paragraph()
        #                 paragraph.add_run('Overall Protein Coding Heatmap:').bold = True
        #                 images = convert_from_path(tem_pdf_path)
        #                 tem_pic_path = Heatmap_pic_folder + '/' + str(pic_no) + '.jpg'
        #                 images[0].save(tem_pic_path, 'JPEG')
        #                 pic_no = pic_no + 1
        #                 output_report_doc.add_picture(tem_pic_path, width=Inches(6.42), height=Inches(3.21))
        #             else:
        #                 print("Overall Protein Coding Heatmap cannot be generated")
        #
        #     if "lincRNA" in gene_type_list:
        #         tem_pdf_path = DE_output_folder_path+ "/linc_heatmap.pdf"
        #         if os.path.isfile(tem_pdf_path):
        #             tem_str = "rm " + tem_pdf_path
        #             os.system(tem_str)
        #         # else:
        #         if linc_padj_path != False and linc_FPKM_path != False:
        #             linc_FPKM_df, C_number = add_information_heatmap_all(linc_FPKM_path)
        #
        #             linc_biomarker_FPKM_df_path = DE_output_folder_path + "/linc_biomarker_FPKM_df.txt"
        #             if os.path.isfile(linc_biomarker_FPKM_df_path):
        #                 tem_str = "rm " + linc_biomarker_FPKM_df_path
        #                 os.system(tem_str)
        #
        #
        #
        #             # tem_str = "rm " + ML_path + "/linc_biomarker_FPKM_df.txt"
        #             # os.system(tem_str)
        #             linc_FPKM_df.to_csv("./linc_biomarker_FPKM_df.txt", sep="\t", na_rep='NA')
        #             R_script_path = directory + "/src/Heatmap_from_DEseq2_LncRNA_all.R"
        #             linc_heatmap_all = subprocess.Popen(['Rscript', R_script_path, str(C_number), linc_padj_path,
        #                               "linc_biomarker_FPKM_df.txt", "linc_heatmap.pdf", pvalue_type, pvalue_threshold, linc_min_FPKM_expression])
        #             linc_heatmap_all.communicate()
        #
        #             if os.path.isfile(tem_pdf_path):
        #                 paragraph = output_report_doc.add_paragraph()
        #                 paragraph.add_run('Overall LincRNA Heatmap:').bold = True
        #                 images = convert_from_path(tem_pdf_path)
        #                 tem_pic_path = Heatmap_pic_folder + '/' + str(pic_no) + '.jpg'
        #                 images[0].save(tem_pic_path, 'JPEG')
        #                 pic_no = pic_no + 1
        #                 output_report_doc.add_picture(tem_pic_path, width=Inches(6.42), height=Inches(3.21))
        #             else:
        #                 print("Overall LincRNA Heatmap cannot be generated")
        #
        #     if "CircRNA" in gene_type_list:
        #         tem_pdf_path = DE_output_folder_path+ "/CircRNA_heatmap.pdf"
        #         if os.path.isfile(tem_pdf_path):
        #             tem_str = "rm " + tem_pdf_path
        #             os.system(tem_str)
        #         # else:
        #         if Circ_padj_path != False and Circ_FPKM_path != False:
        #             Circ_FPKM_df, C_number = add_information_heatmap_all(Circ_FPKM_path)
        #
        #             Circ_biomarker_FPKM_df_path = DE_output_folder_path + "/Circ_biomarker_FPKM_df.txt"
        #             if os.path.isfile(Circ_biomarker_FPKM_df_path):
        #                 tem_str = "rm " + Circ_biomarker_FPKM_df_path
        #                 os.system(tem_str)
        #             Circ_FPKM_df.to_csv("./Circ_biomarker_FPKM_df.txt", sep="\t", na_rep='NA')
        #             R_script_path = directory + "/src/Heatmap_from_DEseq2_CircRNA_all.R"
        #             Circ_heatmap_all = subprocess.Popen(['Rscript', R_script_path, str(C_number), Circ_padj_path,
        #                               "Circ_biomarker_FPKM_df.txt", "CircRNA_heatmap.pdf", pvalue_type, pvalue_threshold, Circ_min_FPKM_expression])
        #             Circ_heatmap_all.communicate()
        #
        #             if os.path.isfile(tem_pdf_path):
        #                 paragraph = output_report_doc.add_paragraph()
        #                 paragraph.add_run('Overall CircRNA Heatmap:').bold = True
        #                 images = convert_from_path(tem_pdf_path)
        #                 tem_pic_path = Heatmap_pic_folder + '/' + str(pic_no) + '.jpg'
        #                 images[0].save(tem_pic_path, 'JPEG')
        #                 pic_no = pic_no + 1
        #                 output_report_doc.add_picture(tem_pic_path, width=Inches(6.42), height=Inches(3.21))
        #             else:
        #                 print("Overall CircRNA Heatmap cannot be generated")



        report_dict = {}

        report_dict["Status"] = []

        for model in model_list:
            report_dict[model] = {}
            report_dict[model]['feature_num'] = []
            report_dict[model]['effective_feature_num'] = []
            report_dict[model]['Accuracy'] = []
            report_dict[model]['Balanced_Accuracy'] = []
            report_dict[model]['ROC_AUC'] = []
            report_dict[model]['Precision'] = []
            report_dict[model]['Recall'] = []
            report_dict[model]['Specificity'] = []
            report_dict[model]['F1'] = []
            report_dict[model]['TN'] = []
            report_dict[model]['FP'] = []
            report_dict[model]['FN'] = []
            report_dict[model]['TP'] = []



        for i in range(len(record_df)):
            JobID = str(record_df.iloc[i].at['JobID'])

            if bjobs_str != '':
                if JobID in JobID_list:
                    Status = JobStatus_list[JobID_list.index(JobID)]
                    report_dict["Status"].append(Status)

                    for model in model_list:
                        report_dict[model]['feature_num'].append("/")
                        report_dict[model]['effective_feature_num'].append("/")
                        report_dict[model]['Accuracy'].append("/")
                        report_dict[model]['Balanced_Accuracy'].append("/")
                        report_dict[model]['ROC_AUC'].append("/")
                        report_dict[model]['Precision'].append("/")
                        report_dict[model]['Recall'].append("/")
                        report_dict[model]['Specificity'].append("/")
                        report_dict[model]['F1'].append("/")
                        report_dict[model]['TN'].append("/")
                        report_dict[model]['FP'].append("/")
                        report_dict[model]['FN'].append("/")
                        report_dict[model]['TP'].append("/")

                    continue

            ID = str(record_df.iloc[i].at['ID'])
            current_Folder_path = str(record_df.iloc[i].at['Folder_path'])
            result_done_path = current_Folder_path + "/" + ID + "/done.txt"

            if os.path.isfile(result_done_path):
                report_dict["Status"].append("Finished")

                heatmap_folder_path = current_Folder_path + "/" + ID + "/heatmaps"

                if os.path.isdir(heatmap_folder_path):
                    shutil.rmtree(heatmap_folder_path)
                os.mkdir(heatmap_folder_path)

                for model in model_list:
                    result_model_path = current_Folder_path + "/" + ID + "/" + model + "_result.csv"
                    report_df = pd.read_csv(result_model_path, index_col=None, sep=",")
                    report_dict[model]['Accuracy'].append(report_df.iloc[0].at['Accuracy'])
                    report_dict[model]['Balanced_Accuracy'].append(report_df.iloc[0].at['Balanced_Accuracy'])
                    report_dict[model]['ROC_AUC'].append(report_df.iloc[0].at['ROC_AUC'])
                    report_dict[model]['Precision'].append(report_df.iloc[0].at['Precision'])
                    report_dict[model]['Recall'].append(report_df.iloc[0].at['Recall'])
                    report_dict[model]['Specificity'].append(report_df.iloc[0].at['Specificity'])
                    report_dict[model]['F1'].append(report_df.iloc[0].at['F1'])
                    report_dict[model]['TN'].append(report_df.iloc[0].at['TN'])
                    report_dict[model]['FP'].append(report_df.iloc[0].at['FP'])
                    report_dict[model]['FN'].append(report_df.iloc[0].at['FN'])
                    report_dict[model]['TP'].append(report_df.iloc[0].at['TP'])


                    report_dict[model]['feature_num'].append(report_df.iloc[0].at['feature_num'])
                    if model == "randomForest" or model == "stacking":
                        feature_list_path = current_Folder_path + "/" + ID + "/" + model + "_imp_features.txt"
                        feature_list_df = pd.read_csv(feature_list_path, sep="\t")
                        feature_list = feature_list_df[(feature_list_df["feature_type"] == "M") & (feature_list_df["feature_score"] != 0)]["feature_name"].tolist()
                        report_dict[model]['effective_feature_num'].append(len(feature_list))
                    else:
                        report_dict[model]['effective_feature_num'].append(report_df.iloc[0].at['feature_num'])

                continue

            if not os.path.isfile(result_done_path):
                report_dict["Status"].append("Terminated")
                for model in model_list:
                    report_dict[model]['feature_num'].append("/")
                    report_dict[model]['effective_feature_num'].append("/")
                    report_dict[model]['Accuracy'].append("/")
                    report_dict[model]['Balanced_Accuracy'].append("/")
                    report_dict[model]['ROC_AUC'].append("/")
                    report_dict[model]['Precision'].append("/")
                    report_dict[model]['Recall'].append("/")
                    report_dict[model]['Specificity'].append("/")
                    report_dict[model]['F1'].append("/")
                    report_dict[model]['TN'].append("/")
                    report_dict[model]['FP'].append("/")
                    report_dict[model]['FN'].append("/")
                    report_dict[model]['TP'].append("/")

                if resubmit == True:
                    resubmit_directory = ML_path + "/" + ID
                    os.chdir(resubmit_directory)
                    system_output = os.popen('bsub < ./pipeline_submit.sh').read()
                    p1 = re.compile(r'[<](.*?)[>]', re.S)
                    JobID_new = re.findall(p1, system_output)[0]
                    record_df.loc[i:i, 'JobID'] = JobID_new
                    ori_record_df.loc[i:i, 'JobID'] = JobID_new
                    os.chdir(ML_path)
                continue

        record_df["Status"] = report_dict["Status"]

        for model in model_list:
            record_df[model + '_feature_num'] = report_dict[model]['feature_num']
            record_df[model + '_effective_feature_num'] = report_dict[model]['effective_feature_num']
            record_df[model + '_Accuracy'] = report_dict[model]['Accuracy']
            record_df[model + '_Balanced_Accuracy'] = report_dict[model]['Balanced_Accuracy']
            record_df[model + '_ROC_AUC'] = report_dict[model]['ROC_AUC']
            record_df[model + '_Precision'] = report_dict[model]['Precision']
            record_df[model + '_Recall'] = report_dict[model]['Recall']
            record_df[model + '_Specificity'] = report_dict[model]['Specificity']
            record_df[model + '_F1'] = report_dict[model]['F1']
            record_df[model + '_TN'] = report_dict[model]['TN']
            record_df[model + '_FP'] = report_dict[model]['FP']
            record_df[model + '_FN'] = report_dict[model]['FN']
            record_df[model + '_TP'] = report_dict[model]['TP']


        ori_record_df.to_csv(ML_path + '/record.csv', index = False)

        report_path = Current_Report_folder_path + '/Job_status.csv'

        record_df = record_df[(record_df["High_Corr_Remove"].isin([High_Corr_Remove_threshold]))]
        record_df.to_csv(report_path, index = False)
        # if Heatmap_generation:
        #     output_report_doc.add_page_break()
        #     tem_str = Current_Report_folder_path + "/HeatMap_Summary.doc"
        #     output_report_doc.save(tem_str)

        model_list_str = '["' + '","'.join(str(e) for e in model_list) + '"]'
        gene_type_list_str = '["' + '","'.join(str(e) for e in gene_type_list) + '"]'

        run_Pipeline_para_0 = src_path + "/Result_Organization_CV_MultiProcess_add_microarray.py"
        run_Pipeline_para_1 = Current_Report_folder_path
        run_Pipeline_para_2 = report_path
        run_Pipeline_para_3 = Target_Metric
        run_Pipeline_para_4 = str(Target_range)
        run_Pipeline_para_5 = current_parent_partition_folder_path #directory
        run_Pipeline_para_6 = DE_type
        run_Pipeline_para_7 = str(CV_partition)
        run_Pipeline_para_8 = summary_type
        run_Pipeline_para_9 = model_list_str
        run_Pipeline_para_10 = gene_type_list_str
        run_Pipeline_para_11 = intersection_check
        run_Pipeline_para_12 = src_path
        run_Pipeline = subprocess.Popen(
            ['python', run_Pipeline_para_0, run_Pipeline_para_1, run_Pipeline_para_2, run_Pipeline_para_3,
             run_Pipeline_para_4, run_Pipeline_para_5, run_Pipeline_para_6, run_Pipeline_para_7, run_Pipeline_para_8,
             run_Pipeline_para_9, run_Pipeline_para_10, run_Pipeline_para_11, run_Pipeline_para_12])
        run_Pipeline.communicate()

        report_path_record_list.append(Current_Report_folder_path)

        print("High_Corr_Remove_threshold " + str(High_Corr_Remove_threshold) + " finished")

        print("parent partition " + str(current_partition) + " finished")
        print(" ")


    file_names_list = ["_CV_performance.csv", "_high_frequency_gene.csv", "_selected_models.csv", "_selected_models_stat_between_partitions.csv"]

    result_summary_folder = directory + "/result_summary"
    if os.path.isdir(result_summary_folder):
        shutil.rmtree(result_summary_folder)
    os.mkdir(result_summary_folder)

    for file_name in file_names_list:
        output_dict = {}
        output_dict["dataset_name"] = []
        sample_file = report_path_record_list[0] + "/" + gene_type_list[0] + file_name
        column_name_list = pd.read_csv(sample_file).columns.tolist()
        for column_name in column_name_list:
            output_dict[column_name] = []
        tem_flag = 0
        for record_Report_folder_path in report_path_record_list:
            current_file = record_Report_folder_path + "/" + gene_type_list[0] + file_name
            current_file_df = pd.read_csv(current_file)
            for column_name in column_name_list:
                output_dict[column_name] = output_dict[column_name] + current_file_df[column_name].tolist()

            tem_list = len(current_file_df) * [dataset_name_list[tem_flag]]
            output_dict["dataset_name"] = output_dict["dataset_name"] + tem_list
            tem_flag += 1

        pd.DataFrame(output_dict).to_csv(result_summary_folder + "/summary_" + gene_type_list[0] + file_name, index=False)

def main():
    parser = argparse.ArgumentParser(formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument("--directory", default=None, type=str, required=True)
    parser.add_argument("--biomarker_target_gene_type", choices=["protein_coding", "lincRNA", "microarray"],
                        default="protein_coding", required=True)
    parser.add_argument("--dataset_name", default="Customized_Dataset", type=str, required=False)
    parser.add_argument("--r", type=str, default="False", help="whether resubmit termiated job", required=False )
    parser.add_argument("--HPC", type=str, default="FALSE", required=False)
    args = parser.parse_args()

    def t_or_f(fs):
        ua = str(fs).upper()
        if 'TRUE'.startswith(ua):
            return True
        elif 'FALSE'.startswith(ua):
            return False
        else:
            return True

    # resubmit_requirment = t_or_f(args.r)
    HPC_parallel_opt = t_or_f(args.HPC)

    Jobstatus_Check(directory=args.directory, biomarker_target_gene_type=args.biomarker_target_gene_type,
                                                   dataset_name=args.dataset_name, resubmit=False, HPC_parallel=HPC_parallel_opt)
    # pipeline_step_1(sample_name=args.sample_name, R1_input=args.R1_input, R2_input=args.R2_input, strandness=args.strandness, directory=args.directory)

if __name__ == '__main__':
    if not sys.warnoptions:
        warnings.simplefilter("ignore")
        os.environ["PYTHONWARNINGS"] = "ignore"  # Also affect subprocesses
    main()

