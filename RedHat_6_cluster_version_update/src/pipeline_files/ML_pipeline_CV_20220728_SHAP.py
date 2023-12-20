# import libraries

import pandas as pd
import matplotlib
matplotlib.use('Agg')
import re
from sklearn.model_selection import train_test_split
import numpy as np
import seaborn as sns; sns.set(font_scale=1.2)
import matplotlib.pyplot as plt
from sklearn.model_selection import cross_val_score
from sklearn.model_selection import RandomizedSearchCV
from sklearn.model_selection import GridSearchCV
from scipy import stats
from scipy.stats import norm
from sklearn.preprocessing import StandardScaler
from sklearn.feature_selection import RFE, SelectFromModel
from sklearn.inspection import permutation_importance
from sklearn.naive_bayes import GaussianNB
from sklearn.ensemble import StackingClassifier
from sklearn import svm
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.linear_model import LogisticRegression
from sklearn import tree
from sklearn.naive_bayes import GaussianNB
from sklearn.model_selection import RepeatedStratifiedKFold
from sklearn.model_selection import StratifiedKFold, KFold
from mlxtend.classifier import StackingCVClassifier
import graphviz
from sklearn.ensemble import RandomForestClassifier
from matplotlib import pyplot as plt
import numpy as np
from numpy import mean
from numpy import std
from sklearn.preprocessing import label_binarize
from sklearn.metrics import roc_curve, auc
from sklearn.feature_selection import RFE
from sklearn.model_selection import RandomizedSearchCV
import xgboost as xgb
from xgboost.sklearn import XGBClassifier
from skopt import gp_minimize
from sklearn.naive_bayes import GaussianNB
from sklearn.base import clone
from kneed import KneeLocator
import numpy.matlib
from sklearn.utils import check_X_y, safe_sqr
from sklearn.pipeline import make_pipeline
from mlxtend.feature_selection import ColumnSelector
from itertools import combinations
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import argparse
import sys
import warnings
import time
from sklearn.metrics import balanced_accuracy_score, precision_score, recall_score, f1_score, roc_auc_score, accuracy_score
from sklearn.preprocessing import label_binarize
import matplotlib.pyplot as plt
from sklearn.model_selection import cross_validate, ShuffleSplit, StratifiedShuffleSplit
from sklearn.utils.validation import _num_samples
from sklearn.utils.validation import _deprecate_positional_args
import math
from math import ceil
from sklearn.utils import resample
from sklearn.utils import check_random_state
from sklearn.externals import joblib
import random
from scipy.signal import argrelextrema
# get_ipython().run_line_magic('matplotlib', 'inline')
from sklearn.preprocessing import MinMaxScaler
import shap
# In[57]:

class BootStrapSplit(ShuffleSplit):
    def __init__(self, n_splits=10, test_size=None, train_size=None,
                 random_state=None):
        super().__init__(
            n_splits=n_splits,
            test_size=test_size,
            train_size=train_size,
            random_state=random_state)
        self._default_test_size = 0.1

    def _iter_indices(self, X, y=None, groups=None):
        n_samples = _num_samples(X)
        n_train, n_test = _validate_shuffle_split(
            n_samples, self.test_size, self.train_size,
            default_test_size=self._default_test_size)

        rng = check_random_state(self.random_state)
        for i in range(self.n_splits):
            # random partition
            permutation = rng.permutation(100)
#             ind_test = permutation[:n_test]
#             ind_train = permutation[n_test:(n_test + n_train)]
            ind_test = resample(np.array(range(n_test+n_train)), replace=True, n_samples=n_test, random_state=permutation[0])
            ind_train = resample(np.array(range(n_test+n_train)), replace=True, n_samples=n_train, random_state=permutation[1])
            yield ind_train, ind_test

def _validate_shuffle_split(n_samples, test_size, train_size,
                            default_test_size=None):
    """
    Validation helper to check if the test/test sizes are meaningful wrt to the
    size of the data (n_samples)
    """
    if test_size is None and train_size is None:
        test_size = default_test_size

    test_size_type = np.asarray(test_size).dtype.kind
    train_size_type = np.asarray(train_size).dtype.kind

    if (test_size_type == 'i' and (test_size >= n_samples or test_size <= 0)
       or test_size_type == 'f' and (test_size <= 0 or test_size >= 1)):
        raise ValueError('test_size={0} should be either positive and smaller'
                         ' than the number of samples {1} or a float in the '
                         '(0, 1) range'.format(test_size, n_samples))

    if (train_size_type == 'i' and (train_size >= n_samples or train_size <= 0)
       or train_size_type == 'f' and (train_size <= 0 or train_size >= 1)):
        raise ValueError('train_size={0} should be either positive and smaller'
                         ' than the number of samples {1} or a float in the '
                         '(0, 1) range'.format(train_size, n_samples))

    if train_size is not None and train_size_type not in ('i', 'f'):
        raise ValueError("Invalid value for train_size: {}".format(train_size))
    if test_size is not None and test_size_type not in ('i', 'f'):
        raise ValueError("Invalid value for test_size: {}".format(test_size))

    if (train_size_type == 'f' and test_size_type == 'f' and
            train_size + test_size > 1):
        raise ValueError(
            'The sum of test_size and train_size = {}, should be in the (0, 1)'
            ' range. Reduce test_size and/or train_size.'
            .format(train_size + test_size))

    if test_size_type == 'f':
        n_test = ceil(test_size * n_samples)
    elif test_size_type == 'i':
        n_test = float(test_size)

    if train_size_type == 'f':
        n_train = floor(train_size * n_samples)
    elif train_size_type == 'i':
        n_train = float(train_size)

    if train_size is None:
        n_train = n_samples - n_test
    elif test_size is None:
        n_test = n_samples - n_train

    if n_train + n_test > n_samples:
        raise ValueError('The sum of train_size and test_size = %d, '
                         'should be smaller than the number of '
                         'samples %d. Reduce test_size and/or '
                         'train_size.' % (n_train + n_test, n_samples))

    n_train, n_test = int(n_train), int(n_test)

    if n_train == 0:
        raise ValueError(
            'With n_samples={}, test_size={} and train_size={}, the '
            'resulting train set will be empty. Adjust any of the '
            'aforementioned parameters.'.format(n_samples, test_size,
                                                train_size)
        )

    return n_train, n_test

class Circ_Biomarker:
    inputCSV = None
    
    def __init__(self, inputCSV):
        self.inputCSV = inputCSV
    
    def load_data(self):
        df = pd.read_csv(self.inputCSV,index_col=0, sep="\t")
        cols = [i for i in df.columns if i not in ['y']]
        global X
        X = df[cols]
        global y
        y = pd.Series(df['y'])

        self.X = X
        self.y = y
        return self


# In[58]:


class RFE_revised(RFE):
    def __init__(self, estimator, n_features_to_select=None, step=1, verbose=0, scoringMetric = "roc_auc"):
        self.estimator = estimator
        self.n_features_to_select = n_features_to_select
        self.step = step
        self.verbose = verbose
        self.scoringMetric = scoringMetric

    def _fit(self, X, y, step_score=None):
        # Parameter step_score controls the calculation of self.scores_
        # step_score is not exposed to users
        # and is used when implementing RFECV
        # self.scores_ will not be calculated when calling _fit through fit

        tags = self._get_tags()
        X, y = check_X_y(X, y, "csc", ensure_min_features=2,
                         force_all_finite=not tags.get('allow_nan', True))
        # Initialization
        n_features = X.shape[1]
        if self.n_features_to_select is None:
            n_features_to_select = n_features // 2
        else:
            n_features_to_select = self.n_features_to_select

        if 0.0 < self.step < 1.0:
            step = int(max(1, self.step * n_features))
        else:
            step = int(self.step)
        if step <= 0:
            raise ValueError("Step must be >0")

        support_ = np.ones(n_features, dtype=np.bool)
        ranking_ = np.ones(n_features, dtype=np.int)

        if step_score:
            self.scores_ = []
#         print("!!!!!!")
        # Elimination
        while np.sum(support_) > n_features_to_select:
            # Remaining features
            features = np.arange(n_features)[support_]

            # Rank the remaining features
            estimator = clone(self.estimator)
            if self.verbose > 0:
                print("Fitting estimator with %d features." % np.sum(support_))

            estimator.fit(X[:, features], y)

            # # Get coefs
            # if hasattr(estimator, 'coef_'):
            #     coefs = estimator.coef_
            # else:
            #     coefs = getattr(estimator, 'feature_importances_', None)
            # if coefs is None:
#                 print("using permutation importance")
            coefs = permutation_importance(estimator, X[:, features], y, scoring= self.scoringMetric, n_jobs = -1, n_repeats=100, random_state=7).importances_mean
#                 raise RuntimeError('The classifier does not expose '
#                                    '"coef_" or "feature_importances_" '
#                                    'attributes')

            # Get ranks
            if coefs.ndim > 1:
                ranks = np.argsort(safe_sqr(coefs).sum(axis=0))
            else:
                ranks = np.argsort(safe_sqr(coefs))

            # for sparse case ranks is matrix
            ranks = np.ravel(ranks)

            # Eliminate the worse features
            threshold = min(step, np.sum(support_) - n_features_to_select)

            # Compute step score on the previous selection iteration
            # because 'estimator' must use features
            # that have not been eliminated yet
            if step_score:
                self.scores_.append(step_score(estimator, features))
            support_[features[ranks][:threshold]] = False
            ranking_[np.logical_not(support_)] += 1

        # Set final attributes
        features = np.arange(n_features)[support_]
        self.estimator_ = clone(self.estimator)
        self.estimator_.fit(X[:, features], y)

        # Compute step score when only n_features_to_select features left
        if step_score:
            self.scores_.append(step_score(self.estimator_, features))
        self.n_features_ = support_.sum()
        self.support_ = support_
        self.ranking_ = ranking_

        return self


# In[59]:


def feature_selection(Model, X, y, scoringMetric, output_report,output_report_fig_index,
                      output_report_index, n_repeats_n_repeats=7, Revised_RFE = False, evaluation = "cv", curve_point = "knee",
                      gene_num_tolerance_for_performance_increase=3, maximum_performance_decrease_tolerance=0.015, non_normalize = True, Tree=False):
#     cb = Circ_Biomarker("./Normal_cancer_UpRegulated_ReadCount.txt").load_data()


    label_number = len(y.value_counts())
    # print("Label", label_number)

    if label_number > 2:
        scoringMetric = "roc_auc_ovo"
    else:
        scoringMetric = "roc_auc"


    n_features_list_all = []
    score_list_all = []
    gene_list_all = []
    n_features_list_coarse = []
    score_list_coarse = []
    n_features_list = []
    score_list = []
    gene_list = []

    # scaler = StandardScaler()
    # scaler.fit(cb.X)
    # X_scaled = pd.DataFrame(scaler.transform(cb.X),columns = cb.X.columns)
    X_scaled = X
    
    # for i in range(X_scaled.shape[1]):
    #     if Revised_RFE == False:
    #         selector = RFE(estimator=clone(Model), n_features_to_select=i+1)
    #     if Revised_RFE == True:
    #         selector = RFE_revised(estimator=clone(Model), n_features_to_select=i+1,scoringMetric = scoringMetric)
    #     selector.fit_transform(X_scaled, y)
    #     X_select = X_scaled[X_scaled.columns[selector.get_support(indices=True)]]
    #     cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=7)
    #     scores = cross_val_score(clone(Model), X_select, y, scoring=scoringMetric, cv=cv, n_jobs=-1, error_score='raise')
    #     n_features_list.append(i+1)
    #     score_list.append(mean(scores))
    if non_normalize:
        X_select = X_scaled.copy()
    else:
        scaler = MinMaxScaler()
        scaler.fit(X_scaled)
        X_select = pd.DataFrame(scaler.transform(X_scaled), columns=X_scaled.columns)

    #############################################################################################################
    ############################  New Code
    consist_flag = 0

    update_gene_list = X_select.columns.tolist()
    current_model = clone(Model)

    while consist_flag == 0:
        X_update = X_select[update_gene_list]
        X_summary = shap.kmeans(X_update, 10)
        model = clone(current_model)
        np.random.seed(3)
        model.fit(X_update, y)
        np.random.seed(7)
        if Tree:
            explainer = shap.TreeExplainer(model=model, seed=7)
        else:
            explainer = shap.KernelExplainer(model=model.predict_proba, data=X_summary, seed=7)
        shap_values = explainer.shap_values(X_update)

        # shap.summary_plot(shap_values, X_update, plot_type="bar", show=False)
        # plt.savefig('./'+ target_file +'/pics/' + str(pic_index) + ".jpg", bbox_inches='tight')
        # plt.show()
        # plt.clf()
        # output_report_doc.add_picture('./'+ target_file +'/pics/' + str(pic_index) + ".jpg")
        # pic_index += 1

        feature_names = X_update.columns

        if type(shap_values) == list:
            rf_resultX_0 = pd.DataFrame(shap_values[0], columns=feature_names)

            vals_0 = np.abs(rf_resultX_0.values).mean(0)

            rf_resultX_1 = pd.DataFrame(shap_values[1], columns=feature_names)

            vals_1 = np.abs(rf_resultX_1.values).mean(0)

            vals = vals_0 + vals_1

        else:
            rf_resultX = pd.DataFrame(shap_values, columns=feature_names)

            vals = np.abs(rf_resultX.values).mean(0)


        shap_importance = pd.DataFrame(list(zip(feature_names, vals)),
                                       columns=['col_name', 'feature_importance_vals'])
        update_gene_list = shap_importance[shap_importance["feature_importance_vals"] > 0]["col_name"].tolist()

        # tem_string = str(len(feature_names)) + " to " + str(len(update_gene_list))
        # output_report_doc.add_paragraph(tem_string)

        if len(update_gene_list) == len(feature_names):
            consist_flag = 1


    output_report[output_report_index] = "The Orignial dataset include " + str(len(X_select.columns.tolist())) + " genes"
    output_report_index += 1


    output_report[output_report_index] = "After Shap value consist, there are" + str(len(update_gene_list)) + " genes"
    output_report_index += 1

    length_after_consist = len(update_gene_list)

    if evaluation == "cv":
        cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=3)
    elif evaluation == "headout":
        cv = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=3)
    elif evaluation == "bootstrap":
        cv = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=3)

    scores = cross_val_score(clone(current_model), X_select[update_gene_list], y, scoring=scoringMetric, cv=cv, n_jobs=-1,
                             error_score='raise')


    output_report[output_report_index] = evaluation + " performance is " + str(mean(scores))
    output_report_index += 1




    # if len(update_gene_list) >= 50:
    #     maximum = 50
    #     tem_shap_importance = shap_importance.copy()
    #     keep_gene_list = tem_shap_importance.sort_values(["feature_importance_vals"], ascending=[False])['col_name'].tolist()[0:50]
    #     fifty_update_gene_list = []
    #     for gene in update_gene_list:
    #         if gene in keep_gene_list:
    #             fifty_update_gene_list.append(gene)
    #
    #     update_gene_list = fifty_update_gene_list.copy()
    #
    # else:
    #     maximum = len(update_gene_list)

    high_dimension_data_flag = 0

    if len(update_gene_list) >= 50:

        high_dimension_data_flag = 1

        # segment_point = max(100, X_select.shape[0])

        interval = math.ceil((X_select[update_gene_list].shape[1] * math.pow(X_select[update_gene_list].shape[0], 2)) / math.pow(10, 6))

        if interval < 10:
            interval = 10
        if interval > 100:
            interval = 100

        # output_report[output_report_index] = "segment_point is " + str(segment_point)
        # output_report_index += 1

        output_report[output_report_index] = "interval is " + str(interval)
        output_report_index += 1

        n_features_list_all.append(len(update_gene_list))
        score_list_all.append(mean(scores))
        gene_list_all.append(update_gene_list)

        while len(update_gene_list) - interval >= 50:
            X_update = X_select[update_gene_list]
            X_summary = shap.kmeans(X_update, 10)
            model = clone(current_model)
            np.random.seed(3)
            model.fit(X_update, y)
            np.random.seed(7)
            if Tree:
                explainer = shap.TreeExplainer(model=model, seed=7)
            else:
                explainer = shap.KernelExplainer(model=model.predict_proba, data=X_summary, seed=7)
            shap_values = explainer.shap_values(X_update)

            # shap.summary_plot(shap_values, X_update, plot_type="bar", show=False)
            # plt.savefig('./' + target_file + '/pics/' + str(pic_index) + ".jpg", bbox_inches='tight')
            # plt.show()
            # plt.clf()
            # output_report_doc.add_picture('./' + target_file + '/pics/' + str(pic_index) + ".jpg")
            # pic_index += 1

            # tem_string = str(len(update_gene_list))
            # output_report_doc.add_paragraph(tem_string)

            feature_names = X_update.columns

            if type(shap_values) == list:
                rf_resultX_0 = pd.DataFrame(shap_values[0], columns=feature_names)

                vals_0 = np.abs(rf_resultX_0.values).mean(0)

                rf_resultX_1 = pd.DataFrame(shap_values[1], columns=feature_names)

                vals_1 = np.abs(rf_resultX_1.values).mean(0)

                vals = vals_0 + vals_1

            else:
                rf_resultX = pd.DataFrame(shap_values, columns=feature_names)

                vals = np.abs(rf_resultX.values).mean(0)

            shap_importance = pd.DataFrame(list(zip(feature_names, vals)),
                                           columns=['col_name', 'feature_importance_vals'])
            update_gene_list = shap_importance["col_name"].tolist()

            tem_list = shap_importance.sort_values(["feature_importance_vals"], ascending=[True])['col_name'].tolist()
            for _ in range(interval):
                remove_gene = tem_list[0]
                update_gene_list.remove(remove_gene)
                tem_list.remove(remove_gene)

            if evaluation == "cv":
                cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=3)
            elif evaluation == "headout":
                cv = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=3)
            elif evaluation == "bootstrap":
                cv = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=3)


            scores = cross_val_score(clone(current_model), X_select[update_gene_list], y, scoring=scoringMetric, cv=cv, n_jobs=-1, error_score='raise')
            n_features_list_all.append(len(update_gene_list))
            score_list_all.append(mean(scores))
            gene_list_all.append(update_gene_list)



        output_report[output_report_index] = "coarse-grained selection from " + str(length_after_consist) + " to " + str(len(update_gene_list)) + " genes"
        output_report_index += 1


        output_report[output_report_index] = "fine-grained selection from " + str(len(update_gene_list)) + " to 1 genes"
        output_report_index += 1

    # maximum = len(update_gene_list)


    if evaluation == "cv":
        cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=3)
    elif evaluation == "headout":
        cv = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=3)
    elif evaluation == "bootstrap":
        cv = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=3)

    scores = cross_val_score(clone(current_model), X_select[update_gene_list], y, scoring=scoringMetric, cv=cv,
                             n_jobs=-1, error_score='raise')

    n_features_list.append(len(update_gene_list))
    score_list.append(mean(scores))
    gene_list.append(update_gene_list)

    if high_dimension_data_flag == 1:
        n_features_list_all.append(len(update_gene_list))
        score_list_all.append(mean(scores))
        gene_list_all.append(update_gene_list)



    for _ in range(len(update_gene_list)-1):
        X_update = X_select[update_gene_list]
        X_summary = shap.kmeans(X_update, 10)
        model = clone(current_model)
        np.random.seed(3)
        model.fit(X_update, y)
        np.random.seed(7)
        if Tree:
            explainer = shap.TreeExplainer(model=model, seed=7)
        else:
            explainer = shap.KernelExplainer(model=model.predict_proba, data=X_summary, seed=7)
        shap_values = explainer.shap_values(X_update)

        # shap.summary_plot(shap_values, X_update, plot_type="bar", show=False)
        # plt.savefig('./'+ target_file +'/pics/' + str(pic_index) + ".jpg", bbox_inches='tight')
        # plt.show()
        # plt.clf()
        # output_report_doc.add_picture('./'+ target_file +'/pics/' + str(pic_index) + ".jpg")
        # pic_index += 1

        # tem_string = str(len(update_gene_list))
        # output_report_doc.add_paragraph(tem_string)

        feature_names = X_update.columns

        if type(shap_values) == list:
            rf_resultX_0 = pd.DataFrame(shap_values[0], columns=feature_names)

            vals_0 = np.abs(rf_resultX_0.values).mean(0)

            rf_resultX_1 = pd.DataFrame(shap_values[1], columns=feature_names)

            vals_1 = np.abs(rf_resultX_1.values).mean(0)

            vals = vals_0 + vals_1

        else:
            rf_resultX = pd.DataFrame(shap_values, columns=feature_names)

            vals = np.abs(rf_resultX.values).mean(0)

        shap_importance = pd.DataFrame(list(zip(feature_names, vals)),
                                       columns=['col_name', 'feature_importance_vals'])
        update_gene_list = shap_importance["col_name"].tolist()
        remove_gene = shap_importance.sort_values(["feature_importance_vals"], ascending=[True])['col_name'].tolist()[0]

        update_gene_list.remove(remove_gene)

        if evaluation == "cv":
            cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=3)
        elif evaluation == "headout":
            cv = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=3)
        elif evaluation == "bootstrap":
            cv = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=3)

        scores = cross_val_score(clone(current_model), X_select[update_gene_list], y, scoring=scoringMetric, cv=cv,
                                 n_jobs=-1, error_score='raise')

        n_features_list.append(len(update_gene_list))
        score_list.append(mean(scores))
        gene_list.append(update_gene_list)


        if high_dimension_data_flag == 1:
            n_features_list_all.append(len(update_gene_list))
            score_list_all.append(mean(scores))
            gene_list_all.append(update_gene_list)


    n_features_list.reverse()
    score_list.reverse()
    if high_dimension_data_flag == 1:
        n_features_list_all.reverse()
        score_list_all.reverse()


    if curve_point.upper() == "KNEE":
        ## Method 1
        #     sensitivity = 10
        #     kn = KneeLocator(n_features_list, score_list, S = sensitivity, curve='concave', direction='increasing', interp_method='interp1d')
        #     while kn.knee == None:
        #         sensitivity  = sensitivity - 1
        #         kn = KneeLocator(n_features_list, score_list, S = sensitivity, curve='concave', direction='increasing', interp_method='interp1d')
        #     knee_point = kn.knee

        ## Method 2
        # curve = score_list
        # nPoints = len(curve)
        # allCoord = np.vstack((range(nPoints), curve)).T
        # np.array([range(nPoints), curve])
        # firstPoint = allCoord[0]
        # lineVec = allCoord[-1] - allCoord[0]
        # lineVecNorm = lineVec / np.sqrt(np.sum(lineVec**2))
        # vecFromFirst = allCoord - firstPoint
        # scalarProduct = np.sum(vecFromFirst * np.matlib.repmat(lineVecNorm, nPoints, 1), axis=1)
        # vecFromFirstParallel = np.outer(scalarProduct, lineVecNorm)
        # vecToLine = vecFromFirst - vecFromFirstParallel
        # distToLine = np.sqrt(np.sum(vecToLine ** 2, axis=1))
        # knee_point = np.argmax(distToLine) + 1

        ## Method 3
        # data = np.array(score_list)
        # local_maximum_index_list = argrelextrema(data, np.greater_equal, order=1)[0].tolist()
        #
        # if len(local_maximum_index_list) > 0:
        #     maximum_performance = score_list[np.argmax(score_list)]
        #     minimum_accepted_performance = maximum_performance
        #     best_point_index = np.argmax(score_list)
        #
        #     while minimum_accepted_performance >= (maximum_performance - maximum_performance_decrease_tolerance):
        #         for tem_index in local_maximum_index_list:
        #             if score_list[tem_index] >= minimum_accepted_performance:
        #                 candidate_best_point_index = tem_index
        #                 break
        #
        #         if best_point_index - candidate_best_point_index >= gene_num_tolerance_for_performance_increase:
        #             best_point_index = candidate_best_point_index
        #
        #         minimum_accepted_performance = minimum_accepted_performance - 0.001
        #
        #     knee_point = best_point_index + 1
        # else:
        #     knee_point = np.argmax(score_list) + 1
        ## Method 4
        data = np.array(score_list)
        local_maximum_index_list = argrelextrema(data, np.greater_equal, order=1)[0].tolist()

        if len(local_maximum_index_list) > 0:
            maximum_performance = score_list[np.argmax(score_list)]

            for tem_index in local_maximum_index_list:
                if score_list[tem_index] >= (maximum_performance - maximum_performance_decrease_tolerance):
                    break
            best_point_index = tem_index
            knee_point = best_point_index + 1
        else:
            knee_point = np.argmax(score_list) + 1
    #     print(idxOfBestPoint)
    elif curve_point.upper() == "HIGHEST":
        knee_point = np.argmax(score_list) + 1


    if high_dimension_data_flag == 1:
        plt.xlabel("n_features group (coarse grained) & individual (fine-grained) feature selection")
        plt.ylabel(scoringMetric)
        plt.plot(n_features_list_all, score_list_all, 'bx-')
        plt.vlines(knee_point, plt.ylim()[0], plt.ylim()[1], linestyles='dashed')

        save_path = "./pics/" + str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches='tight')
        plt.show()
        plt.clf()

    plt.xlabel("n_features individual (fine-grained) feature selection")
    plt.ylabel(scoringMetric)
    plt.plot(n_features_list, score_list, 'bx-')
    plt.vlines(knee_point, plt.ylim()[0], plt.ylim()[1], linestyles='dashed')

    save_path = "./pics/" + str(output_report_index) + '.png'
    output_report[output_report_index] = save_path
    output_report_fig_index.append(output_report_index)
    output_report_index += 1
    plt.savefig(save_path, bbox_inches='tight')
    plt.show()
    plt.clf()

    n_features_list.reverse()
    knee_point_index = n_features_list.index(knee_point)
    knee_point_gene_list = gene_list[knee_point_index]


    # output_report[output_report_index] = "[" + ",".join([str(elem) for elem in score_list]) + "]"
    # output_report_index += 1

    output_report[output_report_index] = "Selected " + str(knee_point) + " Features"
    output_report_index += 1

    output_report[output_report_index] = "[" + ",".join([str(elem) for elem in knee_point_gene_list]) + "]"
    output_report_index += 1

    X_select = X_scaled[knee_point_gene_list]

    selected_columns_index = []

    for gene_name in X_scaled.columns.tolist():
        if gene_name in knee_point_gene_list:
            selected_columns_index.append(True)
        else:
            selected_columns_index.append(False)

    selected_columns_index = np.array(selected_columns_index)


            #############################################################################################################
    ############################  Original Code
#     if X_scaled.shape[1] <= max(100, X_scaled.shape[0]):
#
#         X_select_copy = X_select.copy()
#
#         for i in range(X_scaled.shape[1]):
#             time_start = time.time()
#             if Revised_RFE == False:
#                 selector = RFE(estimator=clone(Model), n_features_to_select=X_scaled.shape[1] - i)
#             if Revised_RFE == True:
#                 selector = RFE_revised(estimator=clone(Model), n_features_to_select=X_scaled.shape[1] - i,
#                                        scoringMetric=scoringMetric)
#             selector.fit_transform(X_select, y)
#             X_select = X_select[X_select.columns[selector.get_support(indices=True)]]
#             # cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=7)
#
#             if evaluation == "cv":
#                 cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=3)
#             elif evaluation == "headout":
#                 cv = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=3)
#             elif evaluation == "bootstrap":
#                 cv = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=3)
#
#             scores = cross_val_score(clone(Model), X_select, y, scoring=scoringMetric, cv=cv, n_jobs=-1,
#                                      error_score='raise')
#             n_features_list.append(X_scaled.shape[1] - i)
#             score_list.append(mean(scores))
#             time_end = time.time()
#             print(i, ',time cost', time_end - time_start, 's')
#
#         n_features_list.reverse()
#         score_list.reverse()
#
#         if curve_point.upper() == "KNEE":
#             ## Method 1
#             #     sensitivity = 10
#             #     kn = KneeLocator(n_features_list, score_list, S = sensitivity, curve='concave', direction='increasing', interp_method='interp1d')
#             #     while kn.knee == None:
#             #         sensitivity  = sensitivity - 1
#             #         kn = KneeLocator(n_features_list, score_list, S = sensitivity, curve='concave', direction='increasing', interp_method='interp1d')
#             #     knee_point = kn.knee
#
#             ## Method 2
#             # curve = score_list
#             # nPoints = len(curve)
#             # allCoord = np.vstack((range(nPoints), curve)).T
#             # np.array([range(nPoints), curve])
#             # firstPoint = allCoord[0]
#             # lineVec = allCoord[-1] - allCoord[0]
#             # lineVecNorm = lineVec / np.sqrt(np.sum(lineVec ** 2))
#             # vecFromFirst = allCoord - firstPoint
#             # scalarProduct = np.sum(vecFromFirst * np.matlib.repmat(lineVecNorm, nPoints, 1), axis=1)
#             # vecFromFirstParallel = np.outer(scalarProduct, lineVecNorm)
#             # vecToLine = vecFromFirst - vecFromFirstParallel
#             # distToLine = np.sqrt(np.sum(vecToLine ** 2, axis=1))
#             # knee_point = np.argmax(distToLine) + 1
#
#             ## Method 3
#             # data = np.array(score_list)
#             # local_maximum_index_list = argrelextrema(data, np.greater_equal, order=1)[0].tolist()
#             #
#             # if len(local_maximum_index_list) > 0:
#             #     maximum_performance = score_list[np.argmax(score_list)]
#             #     minimum_accepted_performance = maximum_performance
#             #     best_point_index = np.argmax(score_list)
#             #
#             #     while minimum_accepted_performance >= (maximum_performance - maximum_performance_decrease_tolerance):
#             #         for tem_index in local_maximum_index_list:
#             #             if score_list[tem_index] >= minimum_accepted_performance:
#             #                 candidate_best_point_index = tem_index
#             #                 break
#             #
#             #         if best_point_index - candidate_best_point_index >= gene_num_tolerance_for_performance_increase:
#             #             best_point_index = candidate_best_point_index
#             #
#             #         minimum_accepted_performance = minimum_accepted_performance - 0.001
#             #
#             #     knee_point = best_point_index + 1
#             # else:
#             #     knee_point = np.argmax(score_list) + 1
#             ## Method 4
#             data = np.array(score_list)
#             local_maximum_index_list = argrelextrema(data, np.greater_equal, order=1)[0].tolist()
#
#             if len(local_maximum_index_list) > 0:
#                 maximum_performance = score_list[np.argmax(score_list)]
#
#                 for tem_index in local_maximum_index_list:
#                     if score_list[tem_index] >= (maximum_performance - maximum_performance_decrease_tolerance):
#                         break
#                 best_point_index = tem_index
#                 knee_point = best_point_index + 1
#             else:
#                 knee_point = np.argmax(score_list) + 1
#
#
#
#
#
#
#
#         #     print(idxOfBestPoint)
#         elif curve_point.upper() == "HIGHEST":
#             knee_point = np.argmax(score_list) + 1
#
#         plt.xlabel("n_features")
#         plt.ylabel(scoringMetric)
#         plt.plot(n_features_list, score_list, 'bx-')
#         plt.vlines(knee_point, plt.ylim()[0], plt.ylim()[1], linestyles='dashed')
#
#         save_path = "./pics/" + str(output_report_index) + '.png'
#         output_report[output_report_index] = save_path
#         output_report_fig_index.append(output_report_index)
#         output_report_index += 1
#         plt.savefig(save_path, bbox_inches = 'tight')
#         plt.show()
#         plt.clf()
#         #     print("11111111111",kn.knee)
#         if Revised_RFE == False:
#             selector = RFE(estimator=clone(Model), n_features_to_select=knee_point)
#         if Revised_RFE == True:
#             selector = RFE_revised(estimator=clone(Model), n_features_to_select=knee_point, scoringMetric=scoringMetric)
#         #     selector = RFE(estimator=clone(Model), n_features_to_select=knee_point)
#
#         # scaler = MinMaxScaler()
#         # scaler.fit(X_scaled)
#         # X_select_copy = pd.DataFrame(scaler.transform(X_scaled), columns=X_scaled.columns)
#
#         selector.fit_transform(X_select_copy, y)
#         selected_columns_index = selector.get_support(indices=True)
#         X_select = X_scaled[X_select_copy.columns[selector.get_support(indices=True)]]
#         print("Selected ", knee_point, " Features")
#
#         output_report[output_report_index] = "[" + ",".join([str(elem) for elem in score_list]) + "]"
#         output_report_index += 1
#
#         output_report[output_report_index] = "Selected " + str(knee_point) + " Features"
#         output_report_index += 1
#
#
#
# ######################################
#
#     else:  # Total Feature number >= 100
#         segment_point = max(100, X_scaled.shape[0])
#
#         interval = math.ceil((X_scaled.shape[1] * math.pow(X_scaled.shape[0], 2)) / math.pow(10, 6))
#
#         if interval < 10:
#             interval = 10
#         if interval > 100:
#             interval = 100
#
#         output_report[output_report_index] = "segment_point is " + str(segment_point)
#         output_report_index += 1
#
#         output_report[output_report_index] = "interval is " + str(interval)
#         output_report_index += 1
#
#         current_point = X_scaled.shape[1]
#
#         # for i in range(X_scaled.shape[1]- segment_point + 1):
#         for i in range(X_scaled.shape[1]):
#             if X_scaled.shape[1] - i > current_point:
#                 continue
#
#             time_start = time.time()
#             if Revised_RFE == False:
#                 selector = RFE(estimator=clone(Model), n_features_to_select=X_scaled.shape[1] - i, step=interval)
#             if Revised_RFE == True:
#                 selector = RFE_revised(estimator=clone(Model), n_features_to_select=X_scaled.shape[1] - i,
#                                        scoringMetric=scoringMetric, step=interval)
#             selector.fit_transform(X_select, y)
#             X_select = X_select[X_select.columns[selector.get_support(indices=True)]]
#             # cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=7)
#
#             if evaluation == "cv":
#                 cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=3)
#             elif evaluation == "headout":
#                 cv = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=3)
#             elif evaluation == "bootstrap":
#                 cv = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=3)
#
#
#             scores = cross_val_score(clone(Model), X_select, y, scoring=scoringMetric, cv=cv, n_jobs=-1, error_score='raise')
#             n_features_list_all.append(X_scaled.shape[1]-i)
#             score_list_all.append(mean(scores))
#             # n_features_list_coarse.append(X_scaled.shape[1]-i)
#             # score_list_coarse.append(mean(scores))
#
#             time_end=time.time()
#             print(i,',time cost',time_end-time_start,'s')
#             current_point = current_point - interval
#
#             if current_point < segment_point:
#                 break
#
#         output_report[output_report_index] = "coarse-grained selection from " + str(X_scaled.shape[1]) + " to " + str(X_scaled.shape[1] - i) + " features"
#         output_report_index += 1
#
#
#         output_report[output_report_index] = "fine-grained selection from " + str( X_scaled.shape[1] - i) + " to 1 feature"
#         output_report_index += 1
#
#         last_i = i
#         last_X_select = X_select.copy()
#
#         for i in range(last_i + 1, X_scaled.shape[1]):
#             time_start = time.time()
#             if Revised_RFE == False:
#                 selector = RFE(estimator=clone(Model), n_features_to_select=X_scaled.shape[1]-i)
#             if Revised_RFE == True:
#                 selector = RFE_revised(estimator=clone(Model), n_features_to_select=X_scaled.shape[1]-i, scoringMetric=scoringMetric)
#             selector.fit_transform(X_select, y)
#             X_select = X_select[X_select.columns[selector.get_support(indices=True)]]
#             # cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=7)
#
#             if evaluation == "cv":
#                 cv = RepeatedStratifiedKFold(n_splits=10, n_repeats=n_repeats_n_repeats, random_state=3)
#             elif evaluation == "headout":
#                 cv = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=3)
#             elif evaluation == "bootstrap":
#                 cv = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=3)
#
#
#
#             scores = cross_val_score(clone(Model), X_select, y, scoring=scoringMetric, cv=cv, n_jobs=-1, error_score='raise')
#             n_features_list_all.append(X_scaled.shape[1]-i)
#             score_list_all.append(mean(scores))
#             n_features_list.append(X_scaled.shape[1]-i)
#             score_list.append(mean(scores))
#             time_end=time.time()
#             print(i,',time cost',time_end-time_start,'s')
#
#
#
#         n_features_list.reverse()
#         score_list.reverse()
#
#
#         n_features_list_all.reverse()
#         score_list_all.reverse()
#
#         if curve_point.upper() == "KNEE":
#             ## Method 1
#         #     sensitivity = 10
#         #     kn = KneeLocator(n_features_list, score_list, S = sensitivity, curve='concave', direction='increasing', interp_method='interp1d')
#         #     while kn.knee == None:
#         #         sensitivity  = sensitivity - 1
#         #         kn = KneeLocator(n_features_list, score_list, S = sensitivity, curve='concave', direction='increasing', interp_method='interp1d')
#         #     knee_point = kn.knee
#
#             ## Method 2
#             # curve = score_list
#             # nPoints = len(curve)
#             # allCoord = np.vstack((range(nPoints), curve)).T
#             # np.array([range(nPoints), curve])
#             # firstPoint = allCoord[0]
#             # lineVec = allCoord[-1] - allCoord[0]
#             # lineVecNorm = lineVec / np.sqrt(np.sum(lineVec**2))
#             # vecFromFirst = allCoord - firstPoint
#             # scalarProduct = np.sum(vecFromFirst * np.matlib.repmat(lineVecNorm, nPoints, 1), axis=1)
#             # vecFromFirstParallel = np.outer(scalarProduct, lineVecNorm)
#             # vecToLine = vecFromFirst - vecFromFirstParallel
#             # distToLine = np.sqrt(np.sum(vecToLine ** 2, axis=1))
#             # knee_point = np.argmax(distToLine) + 1
#
#             ## Method 3
#             # data = np.array(score_list)
#             # local_maximum_index_list = argrelextrema(data, np.greater_equal, order=1)[0].tolist()
#             #
#             # if len(local_maximum_index_list) > 0:
#             #     maximum_performance = score_list[np.argmax(score_list)]
#             #     minimum_accepted_performance = maximum_performance
#             #     best_point_index = np.argmax(score_list)
#             #
#             #     while minimum_accepted_performance >= (maximum_performance - maximum_performance_decrease_tolerance):
#             #         for tem_index in local_maximum_index_list:
#             #             if score_list[tem_index] >= minimum_accepted_performance:
#             #                 candidate_best_point_index = tem_index
#             #                 break
#             #
#             #         if best_point_index - candidate_best_point_index >= gene_num_tolerance_for_performance_increase:
#             #             best_point_index = candidate_best_point_index
#             #
#             #         minimum_accepted_performance = minimum_accepted_performance - 0.001
#             #
#             #     knee_point = best_point_index + 1
#             # else:
#             #     knee_point = np.argmax(score_list) + 1
#             ## Method 4
#             data = np.array(score_list)
#             local_maximum_index_list = argrelextrema(data, np.greater_equal, order=1)[0].tolist()
#
#             if len(local_maximum_index_list) > 0:
#                 maximum_performance = score_list[np.argmax(score_list)]
#
#                 for tem_index in local_maximum_index_list:
#                     if score_list[tem_index] >= (maximum_performance - maximum_performance_decrease_tolerance):
#                         break
#                 best_point_index = tem_index
#                 knee_point = best_point_index + 1
#             else:
#                 knee_point = np.argmax(score_list) + 1
#         #     print(idxOfBestPoint)
#         elif curve_point.upper() == "HIGHEST":
#             knee_point = np.argmax(score_list) + 1
#
#
#         plt.xlabel("n_features group (coarse grained) & individual (fine-grained) feature selection")
#         plt.ylabel(scoringMetric)
#         plt.plot(n_features_list_all,score_list_all, 'bx-')
#         plt.vlines(knee_point, plt.ylim()[0], plt.ylim()[1], linestyles='dashed')
#
#         save_path = "./pics/" +  str(output_report_index) + '.png'
#         output_report[output_report_index] = save_path
#         output_report_fig_index.append(output_report_index)
#         output_report_index += 1
#         plt.savefig(save_path, bbox_inches = 'tight')
#         plt.show()
#         plt.clf()
#
#         plt.xlabel("n_features individual (fine-grained) feature selection")
#         plt.ylabel(scoringMetric)
#         plt.plot(n_features_list,score_list, 'bx-')
#         plt.vlines(knee_point, plt.ylim()[0], plt.ylim()[1], linestyles='dashed')
#
#         save_path = "./pics/" +  str(output_report_index) + '.png'
#         output_report[output_report_index] = save_path
#         output_report_fig_index.append(output_report_index)
#         output_report_index += 1
#         plt.savefig(save_path, bbox_inches = 'tight')
#         plt.show()
#         plt.clf()
#
#
#     #     print("11111111111",kn.knee)
#         if Revised_RFE == False:
#             selector = RFE(estimator=clone(Model), n_features_to_select=knee_point)
#         if Revised_RFE == True:
#             selector = RFE_revised(estimator=clone(Model), n_features_to_select=knee_point,scoringMetric = scoringMetric)
#     #     selector = RFE(estimator=clone(Model), n_features_to_select=knee_point)
#
#         selector.fit_transform(last_X_select, y)
#         selected_columns_index = selector.get_support(indices=True)
#         X_select = X_scaled[last_X_select.columns[selector.get_support(indices=True)]]
#
#
#         # selector.fit_transform(last_X_select, y)
#         # selected_columns_index = selector.get_support(indices=True)
#         # X_select = X_scaled[last_X_select.columns[selector.get_support(indices=True)]]
#         print("Selected ", knee_point, " Features")
#
#         output_report[output_report_index] = "[" + ",".join([str(elem) for elem in score_list]) + "]"
#         output_report_index += 1
#
#         output_report[output_report_index] = "Selected "+ str(knee_point) + " Features"
#         output_report_index += 1
#     #     print(selector.get_support(indices=True))
#     #     print(X_scaled.columns[selector.get_support(indices=True)])
#     #     print("11111111111")
    #############################################################################################################
    return X_select, selected_columns_index, output_report, output_report_fig_index, output_report_index


def feature_importances_evaluation(current_model, X, y, output_report,output_report_fig_index,output_report_index, Tree=False):
    X_summary = shap.kmeans(X, 10)
    model = clone(current_model)
    np.random.seed(3)
    model.fit(X, y)
    np.random.seed(7)
    if Tree:
        explainer = shap.TreeExplainer(model=model, seed=7)
    else:
        explainer = shap.KernelExplainer(model=model.predict_proba, data=X_summary, seed=7)
    shap_values = explainer.shap_values(X)

    shap.summary_plot(shap_values, X, plot_type="bar", max_display=len(X.columns.tolist()), show=False)

    save_path = "./pics/" +  str(output_report_index) + '.png'
    output_report[output_report_index] = save_path

    output_report_fig_index.append(output_report_index)
    output_report_index += 1
    plt.savefig(save_path, bbox_inches='tight')
    plt.show()
    plt.clf()

    # output_report_doc.add_picture('./' + target_file + '/pics/' + str(pic_index) + ".jpg")
    # pic_index += 1

    feature_names = X.columns

    if type(shap_values) == list:
        rf_resultX_0 = pd.DataFrame(shap_values[0], columns=feature_names)

        vals_0 = np.abs(rf_resultX_0.values).mean(0)

        rf_resultX_1 = pd.DataFrame(shap_values[1], columns=feature_names)

        vals_1 = np.abs(rf_resultX_1.values).mean(0)

        vals = vals_0 + vals_1

    else:
        rf_resultX = pd.DataFrame(shap_values, columns=feature_names)

        vals = np.abs(rf_resultX.values).mean(0)

    shap_importance = pd.DataFrame(list(zip(feature_names, vals)),
                                   columns=['col_name', 'feature_importance_vals'])
    features_names = shap_importance["col_name"].tolist()
    imp = shap_importance['feature_importance_vals'].tolist()
    name_index = []
    for i in range(len(imp)):
        name_index.append(i)

    imp, name_index = zip(*sorted(zip(imp, name_index), reverse=True))
    for j in range(len(imp)):
        print('Feature: %s, Score: %.5f' % (features_names[name_index[j]], imp[j]))
        output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]], imp[j])
        output_report_index += 1


    return name_index, imp, features_names, output_report,output_report_fig_index,output_report_index


# In[60]:


class svmModel:    
#     scoringMetric = "accuracy"
#     iterations = 100
#     best_estimator = None
#     best_params= None
#     best_score = None
    
    def __init__ (self, scoringMetric, iterations, X, y, feature_selection = False, evaluation = "cv", curve_point = "knee", pipeline_type="general"):
        # if scoringMetric == 'F1' or scoringMetric == "f1_macro":
        #     self.scoringMetric = "f1_macro"
        # elif scoringMetric == 'AUC' or scoringMetric == "roc_auc":
        #     self.scoringMetric = "roc_auc"
        # else:
        #     print("no metrics requirment be received, use the default metrics F1")
        #     self.scoringMetric = "f1_macro"
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.X = X
        self.y = y
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type=pipeline_type




    
    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)
        
        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index
        
    def runTest(self):   
        # instantiate estimator


        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1

        svm_model = svm.SVC(random_state=7, probability=True, max_iter = 10000)
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X
        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin feature selection"
        output_report_index += 1

        if self.feature_selection == True:
            print("--------------------------------------------------------------------------------")
            output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
            output_report_index += 1
            print("Feature Selection before Modeling")
            output_report[output_report_index] = "Feature Selection before Modeling"
            output_report_index += 1
            X_scaled, selected_columns_index, output_report, output_report_fig_index, output_report_index = feature_selection(Model = svm_model, Revised_RFE = True,
                                                                             X = X_scaled, y = self.y, scoringMetric = self.scoringMetric,
                                                                             output_report = output_report,output_report_fig_index = output_report_fig_index,
                                                                             output_report_index = output_report_index, evaluation = self.evaluation, curve_point=self.curve_point)
        elif self.feature_selection == False:
            selected_columns_index = np.array(range(self.X.shape[1]))
            #         print(X_scaled)
        # define the parameter values that should be searched

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin, fitting"
        output_report_index += 1

        if self.pipeline_type == "general":
            Cs = [0.001, 0.01, 0.1, 1, 5, 10, 100,1000]
    #         gammas = [0.001, 0.01, 0.1, 0.5, 1, 2, 3, 4, 5]
            kernels = ["linear", "poly", "rbf", "sigmoid"]
            decision_function_shapes = ["ovr"]

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                "C": Cs,
    #             "gamma": gammas,
                "kernel": kernels,
                "decision_function_shape": decision_function_shapes
            }

        elif self.pipeline_type == "fresh":
            # define the parameter values that should be searched
            # Original Cs = [0.001, 0.01, 0.1, 1, 5, 10, 100,1000]
            Cs = [0.001, 0.01, 0.1, 1, 5, 10, 100]
            #         gammas = [0.001, 0.01, 0.1, 0.5, 1, 2, 3, 4, 5]
            kernels = ["linear", "poly", "rbf", "sigmoid"]
            decision_function_shapes = ["ovr"]

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                "C": Cs,
                #             "gamma": gammas,
                "kernel": kernels,
                "decision_function_shape": decision_function_shapes
            }



        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)




        # i_iter controls the number of searches
        rand = GridSearchCV(svm_model, param_distrib, cv = cvf, scoring = self.scoringMetric, n_jobs=-1)


        # fit the grid with data
        output_model = rand.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1


        joblib.dump(output_model, './models/Model_SVM.pkl')


        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1

        ##########################################################################
        ##### New Code

        name_index, imp, features_names, output_report, output_report_fig_index, output_report_index = feature_importances_evaluation(
            current_model=rand.best_estimator_, X=X_scaled, y=y, output_report=output_report,
            output_report_fig_index=output_report_fig_index, output_report_index=output_report_index)

        ##########################################################################
        ##### Original Code


#         features_names = []
#         for col in X_scaled.columns:
#             features_names.append(col)
#
#         # importance = rand.best_estimator_.coef_
#         results = permutation_importance(rand.best_estimator_, X_scaled, y, scoring=self.scoringMetric, n_jobs = -1, n_repeats=100, random_state=7)
#         importance = results.importances_mean
#
#         imp = []
#         name_index = []
#
#         # for i,v in enumerate(importance[0]):
#         #         #     name_index.append(i)
#         #         #     imp.append(v)
#         for i,v in enumerate(importance):
#             name_index.append(i)
#             imp.append(v)
#
#
#
# #         features_names = np.array(features_names).reshape(-1)
# #         output_report,output_report_fig_index,output_report_index = self.f_importances(rand.best_estimator_.coef_.reshape(-1), np.array(features_names).reshape(-1),
# #                            output_report = output_report,output_report_fig_index = output_report_fig_index,
# #                            output_report_index = output_report_index)
#
#         output_report,output_report_fig_index,output_report_index = self.f_importances(np.array(imp).reshape(-1), np.array(features_names).reshape(-1),
#                                                                                       output_report = output_report,output_report_fig_index = output_report_fig_index,
#                                                                                        output_report_index = output_report_index)
#
#
#         imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
#         for j in range(len(imp)):
#             print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
#             output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
#             output_report_index += 1

        ##########################################################################

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        # view the results as pandas DataFrame
        pd.DataFrame(rand.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        # examine the best model
        self.best_estimator = rand.best_estimator_
        self.best_params = rand.best_params_
        self.best_score = rand.best_score_
        self.selected_X = X_scaled
        self.selected_columns_index = selected_columns_index
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index


        print(self.best_estimator)
        print("")
        print(output_report)
        print(output_report_fig_index)
        print(output_report_index)
        return self


# In[61]:


class knnModel:
#     scoringMetric = "accuracy"
#     iterations = 100
#     best_estimator = None
#     best_params= None
#     best_score = None
    
    def __init__ (self, scoringMetric, iterations, X, y, feature_selection = False, evaluation =  "cv", curve_point = "knee", pipeline_type="general"):
        # if scoringMetric == 'F1' or scoringMetric == "f1_macro":
        #     self.scoringMetric = "f1_macro"
        # elif scoringMetric == 'AUC' or scoringMetric == "roc_auc":
        #     self.scoringMetric = "roc_auc"
        # else:
        #     print("no metrics requirment be received, use the default metrics F1")
        #     self.scoringMetric = "f1_macro"
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.X = X
        self.y = y
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type = pipeline_type

   
    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)
        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index

    #run KNN model
    def runTest(self):
        # instantiate estimator
        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1

        knn = KNeighborsClassifier()
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin feature selection"
        output_report_index += 1

        if self.feature_selection == True:
            print("--------------------------------------------------------------------------------")
            output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
            output_report_index += 1
            print("Feature Selection before Modeling")
            output_report[output_report_index] = "Feature Selection before Modeling"
            output_report_index += 1
            X_scaled, selected_columns_index, output_report, output_report_fig_index, output_report_index = feature_selection(Model = knn, X = X_scaled, y = self.y, scoringMetric = self.scoringMetric, Revised_RFE = True,n_repeats_n_repeats=3, 
                                                                 output_report = output_report,output_report_fig_index = output_report_fig_index,
                                                                 output_report_index = output_report_index, evaluation = self.evaluation, curve_point=self.curve_point)



        elif self.feature_selection == False:
            selected_columns_index = np.array(range(self.X.shape[1]))
        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin fitting"
        output_report_index += 1

        if self.pipeline_type == "general":
            # define the parameter values that should be searched
            k_range = list(range(1, 11))
            weight_option = ["uniform", "distance"]
            p_range = list(range(1, 11))

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                "n_neighbors": k_range,
                "weights": weight_option,
                "p": p_range
            }

        elif self.pipeline_type == "fresh":
            # define the parameter values that should be searched
            # k_range = list(range(1, 11)) # Original
            k_range = [3, 5, 7, 9, 11]
            weight_option = ["uniform", "distance"]
            # p_range = list(range(1, 11)) # Original
            p_range = [3, 5, 7, 9, 11]

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                "n_neighbors": k_range,
                "weights": weight_option,
                "p": p_range
            }

        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)

        # knn = KNeighborsClassifier(n_jobs=-1)
        # i_iter controls the number of searches
        rand = GridSearchCV(knn, param_distrib, cv = cvf, scoring = self.scoringMetric, n_jobs=-1)

        # fit the grid with data
        output_model = rand.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1



        joblib.dump(output_model, './models/Model_KNN.pkl')

        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1

        ##########################################################################
        ##### New Code

        name_index, imp, features_names, output_report, output_report_fig_index, output_report_index = feature_importances_evaluation(
            current_model=rand.best_estimator_, X=X_scaled, y=y, output_report=output_report,
            output_report_fig_index=output_report_fig_index, output_report_index=output_report_index)

        ##########################################################################
        ##### Original Code



        # features_names = []
        # for col in X_scaled.columns:
        #     features_names.append(col)
        # results = permutation_importance(rand.best_estimator_, X_scaled, y, scoring=self.scoringMetric, n_jobs = -1, n_repeats=100, random_state=7)
        # importance = results.importances_mean
        # imp = []
        # name_index = []
        # for i,v in enumerate(importance):
        #     name_index.append(i)
        #     imp.append(v)
        #
        # output_report,output_report_fig_index,output_report_index = self.f_importances(np.array(imp).reshape(-1), np.array(features_names).reshape(-1),
        #                                                                               output_report = output_report,output_report_fig_index = output_report_fig_index,
        #                                                                                output_report_index = output_report_index)
        #
        #
        # imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
        # for j in range(len(imp)):
        #     print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
        #     output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
        #     output_report_index += 1

        ##########################################################################
        
        # view the results as pandas DataFrame
        pd.DataFrame(rand.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1


        # examine the best model
        self.best_estimator = rand.best_estimator_
        self.best_params = rand.best_params_
        self.best_score = rand.best_score_
        self.selected_X = X_scaled
        self.selected_columns_index = selected_columns_index
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index
        print("")
        return self


# In[62]:


class logRegModel:
    
    def __init__ (self, scoringMetric, iterations, X, y, feature_selection = False, evaluation =  "cv", curve_point = "knee", pipeline_type="general"):
        # if scoringMetric == 'F1' or scoringMetric == "f1_macro":
        #     self.scoringMetric = "f1_macro"
        # elif scoringMetric == 'AUC' or scoringMetric == "roc_auc":
        #     self.scoringMetric = "roc_auc"
        # else:
        #     print("no metrics requirment be received, use the default metrics F1")
        #     self.scoringMetric = "f1_macro"
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.X = X
        self.y = y
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type = pipeline_type

    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)

        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index
    
    #run tests
    def runTest(self):
        # instantiate the model (using the default parameters)
        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1

        logreg = LogisticRegression(random_state = 7, solver="liblinear", n_jobs=-1)
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin feature selection"
        output_report_index += 1


        if self.feature_selection == True:
            print("--------------------------------------------------------------------------------")
            output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
            output_report_index += 1
            print("Feature Selection before Modeling")
            output_report[output_report_index] = "Feature Selection before Modeling"
            output_report_index += 1
            X_scaled, selected_columns_index, output_report, output_report_fig_index, output_report_index = feature_selection(Model = logreg, X = X_scaled, y = self.y, scoringMetric = self.scoringMetric,
                                                                                                                             output_report = output_report,output_report_fig_index = output_report_fig_index, Revised_RFE = True,
                                                                                                                             output_report_index = output_report_index, evaluation = self.evaluation, curve_point=self.curve_point, non_normalize = True)
        elif self.feature_selection == False:
            selected_columns_index = np.array(range(self.X.shape[1]))

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin fitting"
        output_report_index += 1


        if self.pipeline_type == "general":
            # define the parameter values that should be searched
            Cs = [0.0001,0.001, 0.01, 0.1, 1, 5, 10, 100,1000, 10000]
            penalties=['l2', 'l1']
            solver = ['newton-cg', 'liblinear', 'sag', 'saga']
            class_weight_option = ["balanced", None]
            max_iter_option = [1, 3, 5, 10, 30, 50, 100, 300, 500, 10000]


            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                "C": Cs,
                "penalty": penalties,
                "solver": solver,
                'max_iter': max_iter_option,
                "class_weight": class_weight_option,
                "multi_class" : ['ovr', 'multinomial']
            }

        elif self.pipeline_type == "fresh":
            # define the parameter values that should be searched
            # Cs = [0.0001,0.001, 0.01, 0.1, 1, 5, 10, 100,1000, 10000] # Original
            Cs = [0.001, 0.01, 0.1, 1, 5, 10, 100]
            penalties = ['l2', 'l1']
            solver = ['newton-cg', 'lbfgs', 'liblinear', 'sag', 'saga']
            class_weight_option = ["balanced", None]
            # max_iter_option = [1, 3, 5, 10, 30, 50, 100, 300, 500, 10000] # Original
            max_iter_option = [30, 50, 100, 300, 500]

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                "C": Cs,
                "penalty": penalties,
                "solver": solver,
                'max_iter': max_iter_option,
                "class_weight": class_weight_option,
                "multi_class": ['ovr', 'multinomial']
            }
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)

        # i_iter controls the number of searches
        rand = GridSearchCV(logreg, param_distrib, cv = cvf, scoring = self.scoringMetric, n_jobs=-1)
        
        # fit the grid with data
        output_model = rand.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1



        joblib.dump(output_model, './models/Model_logReg.pkl')

        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1

        ##########################################################################
        ##### New Code

        name_index, imp, features_names, output_report, output_report_fig_index, output_report_index = feature_importances_evaluation(
            current_model=rand.best_estimator_, X=X_scaled, y=y, output_report=output_report,
            output_report_fig_index=output_report_fig_index, output_report_index=output_report_index)

        ##########################################################################
        ##### Original Code

        #
        # features_names = []
        # for col in X_scaled.columns:
        #     features_names.append(col)
        # output_report,output_report_fig_index,output_report_index = self.f_importances(rand.best_estimator_.coef_.reshape(-1), np.array(features_names).reshape(-1),
        #                                                                               output_report = output_report,output_report_fig_index = output_report_fig_index,
        #                                                                                output_report_index = output_report_index)
        # importance = rand.best_estimator_.coef_
        # imp = []
        # name_index = []
        # for i,v in enumerate(importance[0]):
        #     name_index.append(i)
        #     imp.append(v)
        # imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
        # for j in range(len(imp)):
        #     print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
        #     output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
        #     output_report_index += 1

        ##########################################################################

        # view the results as pandas DataFrame
        pd.DataFrame(rand.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1


        # examine the best model
        self.best_estimator = rand.best_estimator_
        self.best_params = rand.best_params_
        self.best_score = rand.best_score_
        self.selected_X = X_scaled
        self.selected_columns_index = selected_columns_index
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index
        print("")
        # explanation from the best model
        logreg_explanation = pd.DataFrame(dict(feature=X_scaled.columns, coef= self.best_estimator.coef_[0])) # rand.best_estimator_.coef_: only for linear kernal
        #print(logreg_explanation[logreg_explanation["coef"]!=0])
        return self


# In[63]:


class decisionTreeModel:
    
    #class initialization process
    def __init__ (self, scoringMetric, iterations, X, y, feature_selection = False, evaluation =  "cv", curve_point = "knee", pipeline_type="general"):
        # if scoringMetric == 'F1' or scoringMetric == "f1_macro":
        #     self.scoringMetric = "f1_macro"
        # elif scoringMetric == 'AUC' or scoringMetric == "roc_auc":
        #     self.scoringMetric = "roc_auc"
        # else:
        #     print("no metrics requirment be received, use the default metrics F1")
        #     self.scoringMetric = "f1_macro"
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.X = X
        self.y = y
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type = pipeline_type

    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)
        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index
    
    #run decision tree models
    def runTest(self):  
        # instantiate the model (using the default parameters)
        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1


        Decision_tree = tree.DecisionTreeClassifier(random_state=7)
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin feature selection"
        output_report_index += 1


        if self.feature_selection == True:
            print("--------------------------------------------------------------------------------")
            output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
            output_report_index += 1
            print("Feature Selection before Modeling")
            output_report[output_report_index] = "Feature Selection before Modeling"
            output_report_index += 1
            X_scaled, selected_columns_index, output_report, output_report_fig_index, output_report_index = feature_selection(Model = Decision_tree, X = X_scaled, y = self.y, scoringMetric = self.scoringMetric,
                                                                                                                             output_report = output_report,output_report_fig_index = output_report_fig_index,
                                                                                                                             output_report_index = output_report_index, evaluation = self.evaluation, curve_point=self.curve_point, Tree=True)
        elif self.feature_selection == False:
            selected_columns_index = np.array(range(self.X.shape[1]))

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin fitting"
        output_report_index += 1


        if self.pipeline_type == "general":
            # define the parameter values that should be searched
            Criterion = ["gini", "entropy"]
            Max_depth = [2, 3, 4, 5, 7, 10, 15, None]#[2, 3, 4, 5, None]
            Min_samples_leaf = [1,2,3,5,10]#[1, 3, 5, 10]

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                "criterion": Criterion,
                "splitter":['best', 'random'],
                "max_depth": Max_depth,
                "min_samples_leaf": Min_samples_leaf,
                "max_leaf_nodes": [3,5,7,10,20,None],
                "max_features": ['sqrt', 'log2', None],
                "class_weight": ["balanced", None],
    #             'min_impurity_decrease':[0.0, 0.1,0.2,0.5],
            }
        elif self.pipeline_type == "fresh":
            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {
                # "criterion": ["gini", "entropy"], # Original
                # "splitter":['best', 'random'], # Original
                # "max_depth": [2, 3, 4, 5, 7, 10, 15, None] #[2, 3, 4, 5, None], # Original
                "max_depth": [3, 5, 7, 9, 11, None],
                "min_samples_leaf": [1, 2, 3, 5, 10],  # [1, 3, 5, 10],
                # "max_leaf_nodes": [3,5,7,10,20,None], # Original
                "max_features": ['sqrt', 'log2', None],
                "class_weight": ["balanced", None],
                #             'min_impurity_decrease':[0.0, 0.1,0.2,0.5],

            }

        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        rand = GridSearchCV(Decision_tree, param_distrib, cv = cvf, scoring = self.scoringMetric, n_jobs=-1)

        # fit the grid with data
        output_model =  rand.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1


        joblib.dump(output_model, './models/Model_Decision_tree.pkl')

        self.X_scaled = X_scaled
        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1

        ##########################################################################
        ##### New Code

        name_index, imp, features_names, output_report, output_report_fig_index, output_report_index = feature_importances_evaluation(
            current_model=rand.best_estimator_, X=X_scaled, y=y, output_report=output_report,
            output_report_fig_index=output_report_fig_index, output_report_index=output_report_index, Tree=True)

        ##########################################################################
        ##### Original Code


        # features_names = []
        # for col in X_scaled.columns:
        #     features_names.append(col)
        # importance = rand.best_estimator_.feature_importances_
        # imp = []
        # name_index = []
        # for i,v in enumerate(importance):
        #     name_index.append(i)
        #     imp.append(v)
        #
        #
        # output_report,output_report_fig_index,output_report_index = self.f_importances(np.array(imp).reshape(-1), np.array(features_names).reshape(-1),
        #                                                                               output_report = output_report,output_report_fig_index = output_report_fig_index,
        #                                                                                output_report_index = output_report_index)
        #
        #
        # imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
        # for j in range(len(imp)):
        #     print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
        #     output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
        #     output_report_index += 1
        ##########################################################################

        # view the results as pandas DataFrame
        pd.DataFrame(rand.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        # examine the best model
        self.best_estimator = rand.best_estimator_
        self.best_params = rand.best_params_
        self.best_score = rand.best_score_
        self.selected_X = X_scaled
        self.selected_columns_index = selected_columns_index
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index
        print("")
        
        
        return self

    def graphVisual(self):
        # visualization
        dot_data = tree.export_graphviz(self.best_estimator, out_file=None, 
                                        feature_names=self.X_scaled.columns,
                                        #class_names=['C', 'T'],
                                        filled=True, rounded=True,  
                                        special_characters=True)  
        graph = graphviz.Source(dot_data)  
        
        return graph


# In[64]:


class randomForestModel:
    
    #class initialization process
    def __init__ (self, scoringMetric, iterations, X, y, feature_selection = False, evaluation =  "cv", curve_point = "knee", pipeline_type="general"):
        # if scoringMetric == 'F1' or scoringMetric == "f1_macro":
        #     self.scoringMetric = "f1_macro"
        # elif scoringMetric == 'AUC' or scoringMetric == "roc_auc":
        #     self.scoringMetric = "roc_auc"
        # else:
        #     print("no metrics requirment be received, use the default metrics F1")
        #     self.scoringMetric = "f1_macro"
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.X = X
        self.y = y
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type = pipeline_type
        
    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)
        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index
    #run rforest model
    def runTest(self): 
        # define the parameter values that should be searched
        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1


        Random_forest = RandomForestClassifier(random_state=7)
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin feature selection"
        output_report_index += 1


        if self.feature_selection == True:
            print("--------------------------------------------------------------------------------")
            output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
            output_report_index += 1
            print("Feature Selection before Modeling")
            output_report[output_report_index] = "Feature Selection before Modeling"
            output_report_index += 1
            X_scaled, selected_columns_index, output_report, output_report_fig_index, output_report_index = feature_selection(Model = Random_forest, X = X_scaled, 
                                                                                                                              y = self.y, scoringMetric = self.scoringMetric,n_repeats_n_repeats=3,
                                                                                                                             output_report = output_report,output_report_fig_index = output_report_fig_index,
                                                                                                                             output_report_index = output_report_index, evaluation = self.evaluation, curve_point=self.curve_point, Tree=True)
        elif self.feature_selection == False:
            selected_columns_index = np.array(range(self.X.shape[1]))

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin fitting"
        output_report_index += 1

        if self.pipeline_type == "general":
            # Number of trees in random forest
            n_estimators = [10, 50, 100, 300]
            max_depth = [3, 7, 10, 20, 35, None]
            max_features = ['sqrt', 'log2', None]

            min_samples_leaf = [1, 3, 5, 10, 15]
            min_samples_split = [2, 5, 10, 20]

            oob_score_option = [True, False]

            Criterion = ["gini", "entropy"]
    #         bootstrap = [True, False]

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {'n_estimators': n_estimators,
                           'max_features': max_features,
                           'max_depth': max_depth,
                           'min_samples_split': min_samples_split,
                           'min_samples_leaf': min_samples_leaf,
                             "criterion": Criterion,
                             'oob_score':oob_score_option
    #                        'bootstrap': bootstrap
                            }
        elif self.pipeline_type == "fresh":
            # Number of trees in random forest
            n_estimators = [10, 50, 100, 300]
            # max_depth = [3, 7, 10, 20, 35, None] # Original
            max_features = ['sqrt', 'log2', None]

            min_samples_leaf = [1, 3, 5, 10, 15]
            # min_samples_split = [2, 5, 10, 20] # Original

            # oob_score_option = [True, False] # Original

            # Criterion = ["gini", "entropy"] # Original
            #         bootstrap = [True, False]

            # create a parameter grid: map the parameter names to the values that should be searched
            param_distrib = {'n_estimators': n_estimators,
                             'max_features': max_features,
                             # 'max_depth': max_depth, # Original
                             # 'min_samples_split': min_samples_split, # Original
                             'min_samples_leaf': min_samples_leaf,
                             # "criterion": Criterion, # Original
                             # 'oob_score':oob_score_option # Original
                             #                        'bootstrap': bootstrap
                             }


        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=3, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        # i_iter controls the number of searches
        rand = GridSearchCV(Random_forest, param_distrib, cv = cvf, scoring = self.scoringMetric,verbose=0, n_jobs=-1)

        # fit the grid with data
        output_model = rand.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1


        joblib.dump(output_model, './models/Model_RandomForest.pkl')
        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1

        ##########################################################################
        ##### New Code

        name_index, imp, features_names, output_report, output_report_fig_index, output_report_index = feature_importances_evaluation(
            current_model=rand.best_estimator_, X=X_scaled, y=y, output_report=output_report,
            output_report_fig_index=output_report_fig_index, output_report_index=output_report_index, Tree=True)

        ##########################################################################
        ##### Original Code

        #
        # features_names = []
        # for col in X_scaled.columns:
        #     features_names.append(col)
        # importance = rand.best_estimator_.feature_importances_
        # imp = []
        # name_index = []
        # imp_fig = []
        # selected_features_names_fig = []
        # for i,v in enumerate(importance):
        #     name_index.append(i)
        #     imp.append(v)
        #     if v != 0:
        #         imp_fig.append(v)
        #         selected_features_names_fig.append(features_names[i])
        #
        # if len(imp_fig) != 0 and len(selected_features_names_fig) != 0:
        #
        #     output_report,output_report_fig_index,output_report_index = self.f_importances(np.array(imp_fig).reshape(-1), np.array(selected_features_names_fig).reshape(-1),
        #                        output_report = output_report,output_report_fig_index = output_report_fig_index,
        #                        output_report_index = output_report_index)
        #
        # imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
        # for j in range(len(imp)):
        #     if imp[j] != 0:
        #         print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
        #         output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
        #         output_report_index += 1

        ##########################################################################


        # view the results as pandas DataFrame
        pd.DataFrame(rand.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        # examine the best model
        self.best_estimator = rand.best_estimator_
        self.best_params = rand.best_params_
        self.best_score = rand.best_score_
        self.selected_X = X_scaled
        self.selected_columns_index = selected_columns_index
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index
        print("")

#             print('Feature: %s, Score: %.5f' % (features_names[i],v))
        # plot feature importance
#         plt.bar([x for x in range(len(importance))], importance)
#         plt.show()
        return self


# In[65]:


class GaussianNBModel:
    
    #class initialization process
    def __init__ (self, scoringMetric, iterations, X, y, feature_selection = False, evaluation =  "cv", curve_point = "knee", pipeline_type="general"):
        # if scoringMetric == 'F1' or scoringMetric == "f1_macro":
        #     self.scoringMetric = "f1_macro"
        # elif scoringMetric == 'AUC' or scoringMetric == "roc_auc":
        #     self.scoringMetric = "roc_auc"
        # else:
        #     print("no metrics requirment be received, use the default metrics F1")
        #     self.scoringMetric = "f1_macro"
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.X = X
        self.y = y
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type = pipeline_type
        
    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)
        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index
    
    #run rforest model
    def runTest(self): 
        # define the parameter values that should be searched
        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1

        Bayes = GaussianNB()
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin feature selection"
        output_report_index += 1


        if self.feature_selection == True:
            print("--------------------------------------------------------------------------------")
            output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
            output_report_index += 1
            print("Feature Selection before Modeling")
            output_report[output_report_index] = "Feature Selection before Modeling"
            output_report_index += 1
            X_scaled, selected_columns_index, output_report, output_report_fig_index, output_report_index = feature_selection(Model = Bayes, X = X_scaled, y = self.y, scoringMetric = self.scoringMetric,Revised_RFE = True,
                                                                             output_report = output_report,output_report_fig_index = output_report_fig_index,
                                                                             output_report_index = output_report_index, evaluation = self.evaluation, curve_point=self.curve_point)
        elif self.feature_selection == False:
            selected_columns_index = np.array(range(self.X.shape[1]))
        # create a parameter grid: map the parameter names to the values that should be searched

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin fitting"
        output_report_index += 1

        param_distrib = { } 
        
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        # i_iter controls the number of searches
        rand = GridSearchCV(Bayes, param_distrib, cv = cvf, scoring = self.scoringMetric, n_jobs=-1)

        # fit the grid with data
        output_model = rand.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1

        joblib.dump(output_model, './models/Model_Bayes.pkl')

        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1

        ##########################################################################
        ##### New Code

        name_index, imp, features_names, output_report, output_report_fig_index, output_report_index = feature_importances_evaluation(
            current_model=rand.best_estimator_, X=X_scaled, y=y, output_report=output_report,
            output_report_fig_index=output_report_fig_index, output_report_index=output_report_index)

        ##########################################################################
        ##### Original Code

#         features_names = []
#         for col in X_scaled.columns:
#             features_names.append(col)
#         results = permutation_importance(rand.best_estimator_, X_scaled, y, scoring=self.scoringMetric, n_jobs = -1, n_repeats=100, random_state=7)
#         importance = results.importances_mean
#         imp = []
#         name_index = []
#         for i,v in enumerate(importance):
#             name_index.append(i)
#             imp.append(v)
#
#         output_report,output_report_fig_index,output_report_index = self.f_importances(np.array(imp).reshape(-1), np.array(features_names).reshape(-1),
#                            output_report = output_report,output_report_fig_index = output_report_fig_index,
#                            output_report_index = output_report_index)
#
#
#         imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
#         for j in range(len(imp)):
#             print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
#             output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
#             output_report_index += 1
# #         # plot feature importance
# #         plt.bar([x for x in range(len(importance))], importance)
# #         plt.show()
        ##########################################################################
        # view the results as pandas DataFrame
        pd.DataFrame(rand.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1


        # examine the best model
        self.best_estimator = rand.best_estimator_
        self.best_params = rand.best_params_
        self.best_score = rand.best_score_
        self.selected_X = X_scaled
        self.selected_columns_index = selected_columns_index
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index
        print("")
        return self


# In[66]:


class XGBoostModel:
    #class initialization process
    def __init__ (self, scoringMetric, iterations, X, y, label_number, feature_selection = False, evaluation =  "cv", curve_point = "knee", pipeline_type="general"):
        # if scoringMetric == 'F1' or scoringMetric == "f1_macro":
        #     self.scoringMetric = "f1_macro"
        # elif scoringMetric == 'AUC' or scoringMetric == "roc_auc":
        #     self.scoringMetric = "roc_auc"
        # else:
        #     print("no metrics requirment be received, use the default metrics F1")
        #     self.scoringMetric = "f1_macro"
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.X = X
        self.y = y
        self.label_number = label_number
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type = pipeline_type
        
    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)
        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index
    #run rforest model
    def runTest(self):
        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1

        verbose_num = 0
        if self.label_number > 2:
            objective_para = 'multi:softmax'
        else:
            objective_para = 'binary:logistic'

        eval_metric_para = 'auc'
        n_jobs_option = -1
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X
        param_dist = {'objective':objective_para, 'seed': 7,'silent': 1, 'eval_metric':'auc', 'n_jobs':n_jobs_option}
        xgb_crude_model = xgb.XGBClassifier(**param_dist)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin feature selection"
        output_report_index += 1

        if self.feature_selection == True:
            print("--------------------------------------------------------------------------------")
            output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
            output_report_index += 1
            print("Feature Selection before Modeling")
            output_report[output_report_index] = "Feature Selection before Modeling"
            output_report_index += 1
            X_scaled, selected_columns_index, output_report, output_report_fig_index, output_report_index = feature_selection(Model = xgb_crude_model, X = X_scaled, y = self.y, scoringMetric = self.scoringMetric,
                                                                             output_report = output_report,output_report_fig_index = output_report_fig_index,
                                                                             output_report_index = output_report_index, evaluation = self.evaluation,curve_point=self.curve_point, Tree=True)
        elif self.feature_selection == False:
            selected_columns_index = np.array(range(self.X.shape[1]))

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin fitting"
        output_report_index += 1

        #step 0 n_estimators crude search
        other_params_0 = {'learning_rate': 0.1, 'n_estimators': 500, 'max_depth': 5, 'min_child_weight': 1, 'seed': 7,
                        'subsample': 0.8, 'colsample_bytree': 0.8, 'gamma': 0, 'reg_alpha': 0, 'reg_lambda': 1, 
                         'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        n_estimators_0 = [10, 50, 100, 200, 300, 400, 500, 600, 700, 800,900, 1000]
        param_distrib_0 = {'n_estimators': n_estimators_0} 
        
        model = xgb.XGBClassifier(**other_params_0) 
        
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_0, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['n_estimators'])
        best_n_estimators_0 = optimized_GBM.best_params_['n_estimators']
#         print("=====================================================")
        #step 1 n_estimators
        other_params_1 = {'learning_rate': 0.1, 'n_estimators': 500, 'max_depth': 5, 'min_child_weight': 1, 'seed': 7,
                          'subsample': 0.8, 'colsample_bytree': 0.8, 'gamma': 0, 'reg_alpha': 0, 'reg_lambda': 1, 
                          'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        n_estimators_1 = [ best_n_estimators_0-75, best_n_estimators_0-50, best_n_estimators_0-25, best_n_estimators_0, best_n_estimators_0+25, best_n_estimators_0+50, best_n_estimators_0+75]
        param_distrib_1 = {'n_estimators': n_estimators_1} 
        
        model = xgb.XGBClassifier(**other_params_1)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_1, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['n_estimators'])
        best_n_estimators_1 = optimized_GBM.best_params_['n_estimators']
        
        
        
#         feat_imp = pd.Series(optimized_GBM.best_estimator_.booster().get_fscore()).sort_values(ascending=False)
#         feat_imp.plot(kind='bar', title='Feature Importances')
#         plt.ylabel('Feature Importance Score')
        
        
        
#         print("=====================================================")
        #step 2 max_depth and min_weight
        other_params_2 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': 5, 'min_child_weight': 1, 'seed': 7,
                    'subsample': 0.8, 'colsample_bytree': 0.8, 'gamma': 0, 'reg_alpha': 0, 'reg_lambda': 1, 
                          'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_2 = {'max_depth': [1,2,3, 4, 5, 6, 7, 8, 9, 10], 'min_child_weight': [0.1,0.3,0.5,1, 2, 3, 4, 5, 6]} 
        
        model = xgb.XGBClassifier(**other_params_2)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_2, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['max_depth'])
#         print(optimized_GBM.best_params_['min_child_weight'])
        best_n_max_depth_2 = optimized_GBM.best_params_['max_depth']
        best_n_min_child_weight_2 = optimized_GBM.best_params_['min_child_weight']
        
#         print("=====================================================")
        #step 3 gamma
        other_params_3 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': best_n_max_depth_2, 'min_child_weight': best_n_min_child_weight_2, 'seed': 7,
                    'subsample': 0.8, 'colsample_bytree': 0.8, 'gamma': 0, 'reg_alpha': 0, 'reg_lambda': 1, 
                          'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_3 = {'gamma': [0.0,0.1, 0.2, 0.3, 0.4, 0.5, 0.6,0.7,0.8,0.9,1,5,10,50,100]} 
        
        model = xgb.XGBClassifier(**other_params_3)
        
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_3, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['gamma'])
        best_n_gamma_3 = optimized_GBM.best_params_['gamma']

#         print("=====================================================")
        #step 4 subsample and colsample_bytree crude
        other_params_4 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': best_n_max_depth_2, 'min_child_weight': best_n_min_child_weight_2, 'seed': 7,
                    'subsample': 0.8, 'colsample_bytree': 0.8, 'gamma': best_n_gamma_3, 'reg_alpha': 0, 'reg_lambda': 1, 
                          'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_4 =  {'subsample': [0.5, 0.6, 0.7, 0.8, 0.9,1.0], 'colsample_bytree': [0.5, 0.6, 0.7, 0.8, 0.9,1.0]}
        
        model = xgb.XGBClassifier(**other_params_4)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_4, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['subsample'])
#         print(optimized_GBM.best_params_['colsample_bytree'])
        best_n_subsample_4 = optimized_GBM.best_params_['subsample']
        best_n_colsample_bytree_4 = optimized_GBM.best_params_['colsample_bytree']
        
#         print("=====================================================")
        #step 5 subsample and colsample_bytree 
        other_params_5 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': best_n_max_depth_2, 'min_child_weight': best_n_min_child_weight_2, 'seed': 7,
                    'subsample': 0.8, 'colsample_bytree': 0.8, 'gamma': best_n_gamma_3, 'reg_alpha': 0, 'reg_lambda': 1, 
                          'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_5 =  {'subsample': [best_n_subsample_4-0.05, best_n_subsample_4, best_n_subsample_4+0.05], 
                            'colsample_bytree': [best_n_colsample_bytree_4-0.05, best_n_colsample_bytree_4, best_n_colsample_bytree_4+0.05]}
        
        model = xgb.XGBClassifier(**other_params_5)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_5, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['subsample'])
#         print(optimized_GBM.best_params_['colsample_bytree'])
        best_n_subsample_5 = optimized_GBM.best_params_['subsample']
        best_n_colsample_bytree_5 = optimized_GBM.best_params_['colsample_bytree']

#         print("=====================================================")
        #step 6 reg_alpha and reg_lambda
        other_params_6 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': best_n_max_depth_2, 'min_child_weight': best_n_min_child_weight_2, 'seed': 7,
                    'subsample': best_n_subsample_5, 'colsample_bytree': best_n_colsample_bytree_5, 'gamma': best_n_gamma_3, 
                          'reg_alpha': 0, 'reg_lambda': 1, 'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_6 =  {'reg_alpha': [0, 1e-5,1e-4,1e-3, 1e-2, 0.1, 1, 10, 100], 'reg_lambda': [0, 1e-5,1e-4,1e-3, 1e-2, 0.1, 1, 10, 100]}
        
        model = xgb.XGBClassifier(**other_params_6)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_6, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['reg_alpha'])
#         print(optimized_GBM.best_params_['reg_lambda'])
        best_n_reg_alpha_6 = optimized_GBM.best_params_['reg_alpha']
        best_n_reg_lambda_6 = optimized_GBM.best_params_['reg_lambda']
        reg_alpha_6_unit = best_n_reg_alpha_6
        reg_lambda_6_unit = best_n_reg_lambda_6
        
#         print("=====================================================")
        #step 7 reg_alpha and reg_lambda detail_1
        other_params_7 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': best_n_max_depth_2, 'min_child_weight': best_n_min_child_weight_2, 'seed': 7,
                    'subsample': best_n_subsample_5, 'colsample_bytree': best_n_colsample_bytree_5, 'gamma': best_n_gamma_3, 
                          'reg_alpha': 0, 'reg_lambda': 1, 'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_7 =  {'reg_alpha': [i * best_n_reg_alpha_6 for i in range(1,11)], 
                            'reg_lambda': [i * best_n_reg_lambda_6 for i in range(1,11)]}
        
        model = xgb.XGBClassifier(**other_params_7)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_7, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['reg_alpha'])
#         print(optimized_GBM.best_params_['reg_lambda'])
        best_n_reg_alpha_7 = optimized_GBM.best_params_['reg_alpha']
        best_n_reg_lambda_7 = optimized_GBM.best_params_['reg_lambda']

#         print("=====================================================")
        #step 8 reg_alpha and reg_lambda detail_2
        other_params_8 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': best_n_max_depth_2, 'min_child_weight': best_n_min_child_weight_2, 'seed': 7,
                    'subsample': best_n_subsample_5, 'colsample_bytree': best_n_colsample_bytree_5, 'gamma': best_n_gamma_3, 
                          'reg_alpha': 0, 'reg_lambda': 1, 'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_8 =  {'reg_alpha': [(best_n_reg_alpha_7-(0.5*reg_alpha_6_unit)),best_n_reg_alpha_7,(best_n_reg_alpha_7+(0.5*reg_alpha_6_unit))], 
                            'reg_lambda': [(best_n_reg_lambda_7-(0.5*reg_lambda_6_unit)),best_n_reg_lambda_7,(best_n_reg_lambda_7+(0.5*reg_lambda_6_unit))]}
        
        model = xgb.XGBClassifier(**other_params_8)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_8, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        optimized_GBM.fit(X_scaled, self.y)
        
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['reg_alpha'])
#         print(optimized_GBM.best_params_['reg_lambda'])
        best_n_reg_alpha_8 = optimized_GBM.best_params_['reg_alpha']
        best_n_reg_lambda_8 = optimized_GBM.best_params_['reg_lambda']
        
#         print("=====================================================")
        #step 9 learning rate
        other_params_9 = {'learning_rate': 0.1, 'n_estimators': best_n_estimators_1, 'max_depth': best_n_max_depth_2, 'min_child_weight': best_n_min_child_weight_2, 'seed': 7,
                    'subsample': best_n_subsample_5, 'colsample_bytree': best_n_colsample_bytree_5, 'gamma': best_n_gamma_3, 
                          'reg_alpha': best_n_reg_alpha_8, 'reg_lambda': best_n_reg_lambda_8, 
                          'objective':objective_para,'eval_metric' : eval_metric_para, 'n_jobs':n_jobs_option}
        param_distrib_9 =  {'learning_rate': [0.001,0.01 ,0.03,0.05, 0.07, 0.1, 0.2,0.3,0.5]}
        
        model = xgb.XGBClassifier(**other_params_9)
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)
        optimized_GBM = GridSearchCV(model, param_grid = param_distrib_9, cv = cvf, scoring = self.scoringMetric, verbose=verbose_num, n_jobs=-1)
        output_model = optimized_GBM.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1



        joblib.dump(output_model, './models/Model_XGBoost.pkl')
#         print('参数的最佳取值：{0}'.format(optimized_GBM.best_params_))
#         print('最佳模型得分:{0}'.format(optimized_GBM.best_score_))
#         print(optimized_GBM.best_params_['learning_rate'])
        
#         xgb_param = alg.get_xgb_params()
#         xgtrain = xgb.DMatrix(dtrain[predictors].values, label=dtrain[target].values)
#         cvresult = xgb.cv(xgb_param, xgtrain, num_boost_round=alg.get_params()['n_estimators'], nfold=cv_folds,
#             metrics='auc', early_stopping_rounds=early_stopping_rounds, show_progress=False)
#         alg.set_params(n_estimators=cvresult.shape[0])
        

        print("")
        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1

        ##########################################################################
        ##### New Code

        name_index, imp, features_names, output_report, output_report_fig_index, output_report_index = feature_importances_evaluation(
            current_model=optimized_GBM.best_estimator_, X=X_scaled, y=y, output_report=output_report,
            output_report_fig_index=output_report_fig_index, output_report_index=output_report_index, Tree=True)

        ##########################################################################
        ##### Original Code

        # features_names = []
        # for col in X_scaled.columns:
        #     features_names.append(col)
        # importance = optimized_GBM.best_estimator_.feature_importances_
        # imp = []
        # name_index = []
        # for i,v in enumerate(importance):
        #     name_index.append(i)
        #     imp.append(v)
        #
        # output_report,output_report_fig_index,output_report_index = self.f_importances(np.array(imp).reshape(-1), np.array(features_names).reshape(-1),
        #                    output_report = output_report,output_report_fig_index = output_report_fig_index,
        #                    output_report_index = output_report_index)
        #
        # imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
        # for j in range(len(imp)):
        #     print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
        #     output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
        #     output_report_index += 1

        ##########################################################################

        # view the results as pandas DataFrame
        pd.DataFrame(optimized_GBM.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1


        # examine the best model
        self.best_estimator = optimized_GBM.best_estimator_
        self.best_params = optimized_GBM.best_params_
        self.best_score = optimized_GBM.best_score_
        self.selected_X = X_scaled
        self.selected_columns_index = selected_columns_index    
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index
        
        return self


# In[67]:


class StackingModel:
    
    def __init__ (self, scoringMetric, iterations, X, y, BestEstimator, Selected_columns_index, feature_selection = False, evaluation =  "cv", curve_point = "knee", pipeline_type="general"):
        self.scoringMetric = scoringMetric
        self.iterations = iterations
        self.feature_selection = feature_selection
        self.BestEstimator = BestEstimator
        self.Selected_columns_index = Selected_columns_index
        self.X = X
        self.y = y
        self.evaluation = evaluation
        self.curve_point = curve_point
        self.pipeline_type = pipeline_type
        
    def f_importances(self, coef, names, output_report,output_report_fig_index,output_report_index):
        imp = coef
        imp,names = zip(*sorted(zip(imp,names)))
        plt.barh(range(len(names)), imp, align='center')
        plt.yticks(range(len(names)), names)
#         plt.rcParams['figure.figsize'] = (20.0, 10.0)

        save_path = "./pics/" +  str(output_report_index) + '.png'
        output_report[output_report_index] = save_path
        
        output_report_fig_index.append(output_report_index)
        output_report_index += 1
        plt.savefig(save_path, bbox_inches = 'tight')
        plt.show()
        plt.clf()
        return output_report,output_report_fig_index,output_report_index
    #run tests
    def runTest(self):
        output_report = {}
        output_report_fig_index = []
        output_report_index = 0

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin"
        output_report_index += 1

        # instantiate the model (using the default parameters)
#         scaler = StandardScaler()
#         scaler.fit(X) 
#         X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)
        X_scaled = self.X
#         svm = SVC(random_state=5, kernel='linear', gamma = 0.5, decision_function_shape= 'ovo', C= 0.1,probability=True)
#         knn = KNeighborsClassifier(weights='distance', n_neighbors= 30)
#         logreg = LogisticRegression(random_state = 5, max_iter = 10000, solver='newton-cg', penalty= 'l2', C= 0.1)
        Decision_tree = tree.DecisionTreeClassifier(random_state=7)
#         Random_forest = RandomForestClassifier(random_state=5, n_estimators= 50, min_samples_split= 10, min_samples_leaf=1, max_features= 'sqrt', max_depth= None, bootstrap=True)
        
        lr = LogisticRegression(random_state = 7, max_iter = 10000)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin fitting"
        output_report_index += 1

        pipeline_list = []
        for Estimator_name in self.BestEstimator:
            pipe = make_pipeline(ColumnSelector(cols=(self.Selected_columns_index[Estimator_name]).tolist()),self.BestEstimator[Estimator_name])
            pipeline_list.append(pipe)

        all_comb = []
        for size in range(2, len(self.BestEstimator)+1):
            all_comb += list(combinations(range(len(self.BestEstimator)), r=size))
#         print(all_comb)

        classifiers = []
        for j in range(len(all_comb)):
            tem_list = []
            for k in range(len(all_comb[j])):
                tem_list.append(pipeline_list[all_comb[j][k]])
            tem_tuple = tuple(tem_list)
#             print(tem_tuple)
            classifiers.append(tem_tuple)
#         print(classifiers)

        sclf = StackingCVClassifier(classifiers=Decision_tree,
                            use_probas=True,
                            meta_classifier=lr,
                            random_state=7)

        params = {'meta_classifier__C': [0.001, 0.01, 0.1, 1, 5, 10, 100], #
                   'meta_classifier__solver': ['newton-cg', 'lbfgs', 'liblinear', 'sag', 'saga'],#
                   'classifiers': classifiers}#
    
        # cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        if self.evaluation == "cv":
            cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=7, random_state=7)
        elif self.evaluation == "headout":
            cvf = StratifiedShuffleSplit(n_splits=100, test_size=1 / 3, random_state=7)
        elif self.evaluation == "bootstrap":
            cvf = BootStrapSplit(n_splits=100, test_size=1 / 3, random_state=7)

        grid = GridSearchCV(estimator=sclf, 
                            param_grid=params, 
                            cv=cvf,
                            refit=True, n_jobs=-1)
        output_model = grid.fit(X_scaled, self.y)

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report[output_report_index] = str(random_num) + ". time record begin result organize & saving"
        output_report_index += 1

        joblib.dump(output_model, './models/Model_Stacking.pkl')
        print("--------------------------------------------------------------------------------")
        output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
        output_report_index += 1
        print("Feature Importance from Model")
        output_report[output_report_index] = "Feature Importance from Model"
        output_report_index += 1
        features_names = []
        
        for col in X_scaled.columns:
            features_names.append(col)
        results = permutation_importance(grid.best_estimator_, X_scaled, y, scoring=self.scoringMetric, n_jobs = -1, n_repeats=100, random_state=7)
        importance = results.importances_mean
        imp = []
        name_index = []
        selected_features_names = []

        imp_fig = []
        selected_features_names_fig = []
        for i,v in enumerate(importance):
            # if v > 0:
            name_index.append(i)
            imp.append(v)
            # selected_features_names.append(X.columns[i])
            selected_features_names.append(features_names[i])
            if v != 0:
                imp_fig.append(v)
                # selected_features_names_fig.append(X.columns[i])
                selected_features_names_fig.append(features_names[i])


#         imp = coef
#         imp,names = zip(*sorted(zip(imp,names)))
        if len(imp_fig) != 0 and len(selected_features_names_fig) != 0:
            output_report,output_report_fig_index,output_report_index = self.f_importances(np.array(imp_fig).reshape(-1), np.array(selected_features_names_fig).reshape(-1),
                               output_report = output_report,output_report_fig_index = output_report_fig_index,
                               output_report_index = output_report_index)
        
        imp,name_index = zip(*sorted(zip(imp,name_index), reverse=True))
        for j in range(len(imp)):
            if imp[j] != 0:
                print('Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j]))
                output_report[output_report_index] = 'Feature: %s, Score: %.5f' % (features_names[name_index[j]],imp[j])
                output_report_index += 1
        print("")
        # view the results as pandas DataFrame
        pd.DataFrame(grid.cv_results_)[["params", "mean_test_score", "std_test_score"]]

        time_end = time.time()
        output_report[output_report_index] = str(random_num) + '. time cost: ' + str(time_end - time_start) + 's'
        output_report_index += 1


        # examine the best model
        self.best_estimator = grid.best_estimator_
        self.best_params = grid.best_params_
        self.best_score = grid.best_score_
        self.selected_X = X_scaled
        #         print(list(name_index))
        self.selected_X_after_modling = X_scaled[X_scaled.columns[list(name_index)]]
        self.output_report = output_report
        self.output_report_fig_index = output_report_fig_index
        self.output_report_index = output_report_index
        self.features_names = features_names
        self.imp = imp
        self.name_index = name_index
        # Compute ROC curve and ROC area for each class
#         y_binary = label_binarize(y, classes=[0, 1])
#         y_score = grid.decision_function(X_scaled)
#         fpr,tpr,threshold = roc_curve(y_binary, y_score) ###计算真正率和假正率
#         roc_auc = auc(fpr,tpr) ###计算auc的值
#         print(roc_auc)
#         plt.figure()
#         lw = 2
#         plt.figure(figsize=(10,10))
#         plt.plot(fpr, tpr, color='darkorange', lw=lw, label='ROC curve (area = %0.2f)' % roc_auc) ###假正率为横坐标，真正率为纵坐标做曲线
#         plt.plot([0, 1], [0, 1], color='navy', lw=lw, linestyle='--')
#         plt.xlim([0.0, 1.0])
#         plt.ylim([0.0, 1.05])
#         plt.xlabel('False Positive Rate')
#         plt.ylabel('True Positive Rate')
#         plt.title('Receiver operating characteristic example')
#         plt.legend(loc="lower right")
#         plt.show()

        return self


# In[ ]:


##Correlated Features


# In[51]:


def coorelation_check(X, gene_info_table, threshold = 0.8, linc_padj=False, PC_padj=False, linc_FPKM=False, PC_FPKM=False, Circ_padj=False, Circ_FPKM=False,
                      microarray_FPKM=False, microarray_padj=False, pseudo_foldchange=float(1e-6), pvalue_type="padj"):
    
#     output_report = {}
#     output_report_fig_index = []
#     output_report_index = 0

    # def add_information(padj_path, FPKM_path, CircRNA=False, microarray=False):
    #     if CircRNA:
    #         padj_df = pd.read_csv(padj_path, sep="\t", index_col=0)[["FDR", "PValue"]].rename(columns={"FDR":'padj', "PValue":"pvalue"})
    #     else:
    #         padj_df = pd.read_csv(padj_path, sep="\t", index_col=0)[["padj", "pvalue"]]
    #
    #     FPKM_df = pd.read_csv(FPKM_path, sep="\t", index_col=0)
    #     FPKM_df_transposition = FPKM_df.T
    #     Label_list = []
    #
    #     for i in range(len(FPKM_df_transposition)):
    #         if str(FPKM_df_transposition.index.values[i]).endswith('_C'):
    #             Label_list.append(0)
    #         elif str(FPKM_df_transposition.index.values[i]).endswith('_T'):
    #             Label_list.append(1)
    #         else:
    #             print("Label not match")
    #
    #     if microarray:
    #         column_name_list = FPKM_df_transposition.columns.values.tolist()
    #         for column_name in column_name_list:
    #             FPKM_df_transposition[column_name] = 2**FPKM_df_transposition[column_name]
    #
    #     FPKM_df_transposition['label'] = Label_list
    #     FPKM_df_transposition_0 = FPKM_df_transposition[FPKM_df_transposition['label'].isin([0])]
    #     FPKM_df_transposition_1 = FPKM_df_transposition[FPKM_df_transposition['label'].isin([1])]
    #
    #     FPKM_df_transposition_0_transform = FPKM_df_transposition_0.drop(['label'], axis=1).T
    #     FPKM_df_transposition_1_transform = FPKM_df_transposition_1.drop(['label'], axis=1).T
    #
    #     FPKM_df_transposition_0_transform["mean"] = FPKM_df_transposition_0_transform.mean(axis=1)
    #     FPKM_df_transposition_1_transform["mean"] = FPKM_df_transposition_1_transform.mean(axis=1)
    #
    #     FPKM_df_transposition_0_transform["std"] = FPKM_df_transposition_0_transform.std(axis=1)
    #     FPKM_df_transposition_1_transform["std"] = FPKM_df_transposition_1_transform.std(axis=1)
    #
    #     FPKM_df_C_mean = FPKM_df_transposition_0_transform[["mean"]].rename(columns={'mean': 'mean_Control'})
    #     FPKM_df_T_mean = FPKM_df_transposition_1_transform[["mean"]].rename(columns={'mean': 'mean_Treatment'})
    #
    #     FPKM_df_C_std = FPKM_df_transposition_0_transform[["std"]].rename(columns={'std': 'std_Control'})
    #     FPKM_df_T_std = FPKM_df_transposition_1_transform[["std"]].rename(columns={'std': 'std_Treatment'})
    #
    #     FPKM_df_with_mean = pd.merge(FPKM_df_C_mean, FPKM_df_T_mean, left_index=True, right_index=True)
    #     FPKM_df_with_std = pd.merge(FPKM_df_C_std, FPKM_df_T_std, left_index=True, right_index=True)
    #
    #     FPKM_df_with_all = pd.merge(FPKM_df_with_mean, FPKM_df_with_std, left_index=True, right_index=True)
    #
    #     abs_log2_list = []
    #     foldchange_sd_score_list = []
    #     for i in range(len(FPKM_df_with_all)):
    #         mean_C = FPKM_df_with_all.iloc[i].at['mean_Control']
    #         mean_T = FPKM_df_with_all.iloc[i].at['mean_Treatment']
    #         if mean_C == 0 or mean_T == 0:
    #             mean_C = mean_C + pseudo_foldchange
    #             mean_T = mean_T + pseudo_foldchange
    #         tem_abs_log2_foldchange = abs(math.log(mean_T / mean_C, 2))
    #         abs_log2_list.append(tem_abs_log2_foldchange)
    #
    #         std_C = FPKM_df_with_all.iloc[i].at['std_Control']
    #         std_T = FPKM_df_with_all.iloc[i].at['std_Treatment']
    #         max_std = max(std_C, std_T)
    #         if max_std == 0:
    #             max_std = max_std + pseudo_foldchange
    #             tem_abs_log2_foldchange = tem_abs_log2_foldchange + pseudo_foldchange
    #
    #         tem_foldchange_sd_score = tem_abs_log2_foldchange / max_std
    #         foldchange_sd_score_list.append(tem_foldchange_sd_score)
    #
    #
    #
    #
    #     FPKM_df_with_all["abs(log2(fold-change))"] = abs_log2_list
    #     FPKM_df_with_all["fold-change_std_score"] = foldchange_sd_score_list
    #
    #     info_df = pd.merge(FPKM_df_with_all, padj_df, left_index=True, right_index=True, how='left')
    #
    #     return info_df

    linc_flag = 0
    PC_flag = 0
    Circ_flag = 0
    microarray_flag = 0
    add_info_flag = 0

    if linc_padj != False and linc_FPKM!= False:
        # linc_info_df = add_information(linc_padj, linc_FPKM)
        linc_info_df = gene_info_table.copy()
        linc_flag = 1
        add_info_flag = 1


    if PC_padj != False and PC_FPKM!= False:
        # PC_info_df = add_information(PC_padj, PC_FPKM)
        PC_info_df = gene_info_table.copy()
        PC_flag = 1
        add_info_flag = 1

    if Circ_padj != False and Circ_FPKM!= False:
        # Circ_info_df = add_information(Circ_padj, Circ_FPKM, CircRNA=True)
        Circ_info_df = gene_info_table.copy()
        Circ_flag = 1
        add_info_flag = 1

    if microarray_padj != False and microarray_FPKM!= False:
        # microarray_info_df = add_information(microarray_padj, microarray_FPKM, microarray=True)
        microarray_info_df = gene_info_table.copy()
        microarray_flag = 1
        add_info_flag = 1


    if add_info_flag == 1:
        if linc_flag == 1 and PC_flag == 1:
            info_df = pd.concat([linc_info_df, PC_info_df])
        elif linc_flag == 1 and PC_flag == 0 and Circ_flag == 0 and microarray_flag == 0:
            info_df = linc_info_df
        elif linc_flag == 0 and PC_flag == 1 and Circ_flag == 0 and microarray_flag == 0:
            info_df = PC_info_df
        elif linc_flag == 0 and PC_flag == 0 and Circ_flag == 1 and microarray_flag == 0:
            info_df = Circ_info_df
        elif linc_flag == 0 and PC_flag == 0 and Circ_flag == 0 and microarray_flag == 1:
            info_df = microarray_info_df




    correlated_features = set()
    correlation_matrix = X.corr()
#     features = []
    features_pairs = {'No.':[],'Drop Feature':[],'Keep Feature':[],'Correlation':[]}
    flag = 0

    if add_info_flag == 1:
        for i in range(len(correlation_matrix.columns)):
            for j in range(i):
                if (abs(correlation_matrix.iloc[i, j]) > threshold) and (correlation_matrix.columns[j] not in correlated_features) and (correlation_matrix.columns[i] not in correlated_features):
                    colname_1 = correlation_matrix.columns[i]
                    colname_2 = correlation_matrix.columns[j]
                    if pvalue_type == "padj":
                        # pvalue_1 = info_df[(info_df.index==colname_1)]["padj"].tolist()[0]
                        # pvalue_2 = info_df[(info_df.index == colname_2)]["padj"].tolist()[0]

                        pvalue_1 = float(info_df.at[colname_1, "padj"])
                        pvalue_2 = float(info_df.at[colname_2, "padj"])

                    elif pvalue_type == "pvalue":
                        # pvalue_1 = info_df[(info_df.index==colname_1)]["pvalue"].tolist()[0]
                        # pvalue_2 = info_df[(info_df.index == colname_2)]["pvalue"].tolist()[0]

                        pvalue_1 = float(info_df.at[colname_1, "pvalue"])
                        pvalue_2 = float(info_df.at[colname_2, "pvalue"])

                    # fcstd_Score_1 = info_df[(info_df.index==colname_1)]["fold-change_std_score"].tolist()[0]
                    # fcstd_Score_2 = info_df[(info_df.index == colname_2)]["fold-change_std_score"].tolist()[0]

                    # fcstd_Score_1 = info_df[(info_df.index==colname_1)]["score"].tolist()[0]
                    # fcstd_Score_2 = info_df[(info_df.index == colname_2)]["score"].tolist()[0]

                    fcstd_Score_1 = float(info_df.at[colname_1, "score"])
                    fcstd_Score_2 = float(info_df.at[colname_2, "score"])


                    # max_mean_FPKM_1 = max(info_df[(info_df.index==colname_1)]["mean_Control"].tolist()[0], info_df[(info_df.index==colname_1)]["mean_Treatment"].tolist()[0])
                    # max_mean_FPKM_2 = max(info_df[(info_df.index == colname_2)]["mean_Control"].tolist()[0], info_df[(info_df.index == colname_2)]["mean_Treatment"].tolist()[0])

                    abs_log2_foldchange_1 = info_df[(info_df.index==colname_1)]["abs(log2(fold-change))"].tolist()[0]
                    abs_log2_foldchange_2 = info_df[(info_df.index == colname_2)]["abs(log2(fold-change))"].tolist()[0]

                    abs_log2_foldchange_1 = float(info_df.at[colname_1, "abs(log2(fold-change))"])
                    abs_log2_foldchange_2 = float(info_df.at[colname_2, "abs(log2(fold-change))"])

                    if microarray_flag == 1:
                        if pvalue_1 <= pvalue_2:
                            colname_keep = colname_1
                            colname_drop = colname_2
                        elif pvalue_1 > pvalue_2:
                            colname_keep = colname_2
                            colname_drop = colname_1

                    else:
                        if ((abs_log2_foldchange_1 >= 1) and (abs_log2_foldchange_2 >= 1) and (pvalue_1 <= 0.05) and (pvalue_2 <= 0.05)):
                            if fcstd_Score_1 >= fcstd_Score_2:
                                colname_keep = colname_1
                                colname_drop = colname_2
                            elif fcstd_Score_1 < fcstd_Score_2:
                                colname_keep = colname_2
                                colname_drop = colname_1

                        else:
                            if ((pvalue_1 <= 0.05) and (pvalue_2 <= 0.05)):
                                if ((abs_log2_foldchange_1 < 1) and (abs_log2_foldchange_2 < 1)):
                                    if fcstd_Score_1 >= fcstd_Score_2:
                                        colname_keep = colname_1
                                        colname_drop = colname_2
                                    elif fcstd_Score_1 < fcstd_Score_2:
                                        colname_keep = colname_2
                                        colname_drop = colname_1
                                else:
                                    if abs_log2_foldchange_1 >= abs_log2_foldchange_2:
                                        colname_keep = colname_1
                                        colname_drop = colname_2
                                    elif abs_log2_foldchange_1 < abs_log2_foldchange_2:
                                        colname_keep = colname_2
                                        colname_drop = colname_1
                            else:
                                if pvalue_1 <= pvalue_2:
                                    colname_keep = colname_1
                                    colname_drop = colname_2
                                elif pvalue_1 > pvalue_2:
                                    colname_keep = colname_2
                                    colname_drop = colname_1

                    features_pairs['Drop Feature'].append(colname_drop)
                    features_pairs['Keep Feature'].append(colname_keep)
                    features_pairs['Correlation'].append(abs(correlation_matrix.iloc[i, j]))
                    features_pairs['No.'].append(flag)
                    flag = flag + 1
                    correlated_features.add(colname_drop)





    elif add_info_flag == 0:
        for i in range(len(correlation_matrix.columns)):
            for j in range(i):
                if (abs(correlation_matrix.iloc[i, j]) > threshold) and (correlation_matrix.columns[j] not in correlated_features):
                    colname_1 = correlation_matrix.columns[i]
                    colname_2 = correlation_matrix.columns[j]
    #                 features.append([colname_1,colname_2])
                    features_pairs['Drop Feature'].append(colname_1)
                    features_pairs['Keep Feature'].append(colname_2)
                    features_pairs['Correlation'].append(abs(correlation_matrix.iloc[i, j]))
                    features_pairs['No.'].append(flag)
                    flag = flag + 1
                    correlated_features.add(colname_1)
    updated_X = X
    for feature_name in correlated_features:
        updated_X = updated_X.drop(feature_name, axis=1)
    print("================================================================================")
    print("High Correlation Feature Pairs:")
    print("--------------------------------------------------------------------------------")
    
    
#     output_report[output_report_index] = '======================================================================'
#     output_report_index += 1
#     output_report[output_report_index] = "High Correlation Feature Pairs:"
#     output_report_index += 1
#     output_report[output_report_index] = "--------------------------------------------------------------------------------------------------------------------"
#     output_report_index += 1
    
    df_correlated_features = pd.DataFrame(features_pairs)
    print(df_correlated_features)
    print("================================================================================")
#     output_report[output_report_index] = '======================================================================'
#     output_report_index += 1
    return updated_X, features_pairs, df_correlated_features


# In[68]:


def circBiomarker(input_file,info_table_path, Model=["svm","knn","logReg","decisionTree","randomForest","bayes","xgboost","stacking"], iterations=100,
                  scoringMetric="AUC", feature_selection = True, test_file = False, val_method_selection="cv", val_method_final=["cv", "headout", "bootstrap"], corrthreshold = 0.8, curve_point = "knee", pipeline_type = "general",
                  linc_padj=False, PC_padj=False, linc_FPKM=False, PC_FPKM=False, Circ_padj=False, Circ_FPKM=False, pseudo_foldchange=float(1e-6), pvalue_type="padj", pvalue_threshold = 0.05, foldchange_threshold = 1.0,
                  PC_maxmin_remove_threshold = 1.0,  linc_maxmin_remove_threshold = 0.001, Circ_maxmin_remove_threshold = 0, final_evaluation=False,
                  microarray_padj=False, microarray_FPKM=False, microarray_maxmin_remove_threshold= 0):

    if not os.path.isdir('./pics'):
        os.mkdir('./pics')

    if not os.path.isdir('./models'):
        os.mkdir('./models')
    
    output_report_doc = Document()
    output_report_doc.add_heading(' Model Selection Report ', 0)

    time_start = time.time()
    random_num = random.randint(0, 99)
    output_report_doc.add_paragraph(str(random_num) + ". time record begin preparation ")

    gene_info_table = pd.read_csv(info_table_path, sep="\t", index_col=0)


    cb = Circ_Biomarker(input_file).load_data()
    X = cb.X
    y = cb.y
    label_number = len(cb.y.value_counts())
    print("Label", label_number)
    evaluation = val_method_selection
    evaluation_final = val_method_final
    if scoringMetric == 'F1':
        scoringMetric = "f1_macro"
    elif scoringMetric == 'AUC':
        if label_number > 2:
            scoringMetric = "roc_auc_ovo"
        else:
            scoringMetric = "roc_auc"
    elif scoringMetric == 'BACC':
        scoringMetric = "balanced_accuracy"
    elif scoringMetric == 'ACC':
        scoringMetric = "accuracy"
    else:
        print("no metrics requirment be received, use the default metrics F1")
        output_report_doc.add_paragraph('no metrics requirment be received, use the default metrics F1')
        scoringMetric = "f1_macro"

    BestScore={}
    BestEstimator={}
    Selected_X={}
    Selected_columns_index = {}
    Best_Params = {}
    FI_features_names = {}
    FI_imp = {}
    FI_name_index = {}
    Model_output_report = {}
    Model_output_report_fig_index = {}
    # scaler = StandardScaler()
    # scaler.fit(X)
    # print(X.shape)
    # X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)

    X_scaled = X

    time_end = time.time()
    output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's preparation end')

    time_start = time.time()
    random_num = random.randint(0, 99)
    output_report_doc.add_paragraph(str(random_num) + ". time record begin high correlation remove")

    if feature_selection == True:
        output_report_doc.add_paragraph('======================================================================')
#         output_report_doc.add_paragraph('High Correlation Feature Pairs:')

        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('low expression gene remove:').bold = True

        gene_type_flag = 0

        if PC_padj != False and PC_FPKM != False and linc_padj == False and linc_FPKM == False and Circ_padj == False and Circ_FPKM == False and microarray_padj == False and microarray_FPKM == False:
            gene_type = "Protein Coding Gene"
            remove_threshold = PC_maxmin_remove_threshold
            output_report_doc.add_paragraph('gene type is '+ gene_type)
        elif PC_padj == False and PC_FPKM == False and linc_padj != False and linc_FPKM != False and Circ_padj == False and Circ_FPKM == False and microarray_padj == False and microarray_FPKM == False:
            gene_type = "lincRNA"
            remove_threshold = linc_maxmin_remove_threshold
            output_report_doc.add_paragraph('gene type is ' + gene_type)
        elif PC_padj == False and PC_FPKM == False and linc_padj == False and linc_FPKM == False and Circ_padj != False and Circ_FPKM != False and microarray_padj == False and microarray_FPKM == False:
            gene_type = "Circular RNA"
            remove_threshold = Circ_maxmin_remove_threshold
            output_report_doc.add_paragraph('gene type is ' + gene_type)
        elif PC_padj == False and PC_FPKM == False and linc_padj == False and linc_FPKM == False and Circ_padj == False and Circ_FPKM == False and microarray_padj != False and microarray_FPKM != False:
            gene_type = "Microarray Data"
            remove_threshold = microarray_maxmin_remove_threshold
            output_report_doc.add_paragraph('gene type is ' + gene_type)
        else:
            gene_type_flag = 1
            output_report_doc.add_paragraph('gene type cannot recognize')

        # if gene_type_flag != 1:
        #     low_expression_gene_list = []
        #     gene_col_name_list = X.max().to_frame().index.tolist()
        #     gene_col_max_list = X.max().to_frame()[0].tolist()
        #     for i in range(len(gene_col_max_list)):
        #         if abs(gene_col_max_list[i]) < remove_threshold:
        #             low_expression_gene_list.append(gene_col_name_list[i])
        #
        #     for feature_name in low_expression_gene_list:
        #         X_scaled = X_scaled.drop(feature_name, axis=1)
        #
        #     output_report_doc.add_paragraph('removed follow low expression level features, remove threshold: ' + str(remove_threshold))
        #
        #     COLUMNS = 2
        #     ROWS = len(low_expression_gene_list)
        #
        #     table = output_report_doc.add_table(rows=ROWS + 1, cols=COLUMNS, style='Table Grid')
        #     hdr_cells = table.rows[0].cells
        #     hdr_cells[0].text = 'No.'
        #     hdr_cells[1].text = 'Drop Feature'
        #
        #     list_0 = [i for i in range(1, ROWS + 1)]
        #     list_1 = low_expression_gene_list
        #     table_cells = table._cells
        #
        #     for i in range(1, ROWS + 1):
        #         row_cells = table_cells[i * COLUMNS:(i + 1) * COLUMNS]
        #         # row_cells = table.add_row().cells
        #         row_cells[0].text = str(list_0[i - 1])
        #         row_cells[1].text = str(list_1[i - 1])
        #
        # output_report_doc.add_paragraph('======================================================================')

        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('High Correlation Feature Pairs:').bold = True

        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        X_scaled, features_pairs, df_correlated_features = coorelation_check(X_scaled, threshold=corrthreshold,
                                                                             linc_padj=linc_padj, PC_padj=PC_padj, linc_FPKM=linc_FPKM,
                                                                             PC_FPKM=PC_FPKM, Circ_padj=Circ_padj, Circ_FPKM=Circ_FPKM,
                                                                             microarray_FPKM=microarray_FPKM, microarray_padj=microarray_padj,
                                                                             pseudo_foldchange=pseudo_foldchange, pvalue_type=pvalue_type, gene_info_table=gene_info_table)

        COLUMNS = 4
        ROWS = len(df_correlated_features.index)

        table = output_report_doc.add_table(rows=ROWS+1, cols=COLUMNS,style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No.'
        hdr_cells[1].text = 'Drop Feature'
        hdr_cells[2].text = 'Keep Feature'
        hdr_cells[3].text = 'Correlation'
        list_0=list(df_correlated_features['No.'])
        list_1=list(df_correlated_features['Drop Feature'])
        list_2=list(df_correlated_features['Keep Feature'])
        list_3=list(df_correlated_features['Correlation'])
        table_cells = table._cells

        for i in range(1, ROWS+1):
            row_cells = table_cells[i * COLUMNS:(i + 1) * COLUMNS]
            # row_cells = table.add_row().cells
            row_cells[0].text = str(list_0[i-1])
            row_cells[1].text = str(list_1[i-1])
            row_cells[2].text = str(list_2[i-1])
            row_cells[3].text = str(round(list_3[i-1],3))
        output_report_doc.add_paragraph('======================================================================')
        
    # scaler = StandardScaler()
    # scaler.fit(X)
    # X_scaled = pd.DataFrame(scaler.transform(X),columns = X.columns)

    print(X_scaled.shape)
    print("")

    time_end = time.time()
    output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's high correlation remove end')

    j = 1
    print("Scoring Metric is ", scoringMetric, ", evaluation method is ", evaluation, ", evaluation method for final score is ", evaluation_final)
    tem_str = "Scoring Metric is " + scoringMetric + ", evaluation method is " + evaluation + ", evaluation method for final score is " + ' '.join([str(elem) for elem in evaluation_final])
    output_report_doc.add_paragraph(tem_str)
    if "svm" in Model:
        print("================================================================================")
        print(j,"Running SVM Model")
        output_report_doc.add_paragraph('======================================================================')
#         output_report_doc.add_paragraph('%d. Running SVM Model' % (j))
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running SVM Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin svm total")

        svm = svmModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, feature_selection = feature_selection, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["svm"] = svm.best_score
        BestEstimator["svm"] = svm.best_estimator
        Selected_X["svm"] = svm.selected_X
        Selected_columns_index["svm"] = svm.selected_columns_index
        Best_Params["svm"] = svm.best_params
        FI_features_names["svm"] = svm.features_names
        FI_imp["svm"] = svm.imp
        FI_name_index["svm"] = svm.name_index
        Model_output_report["svm"] = svm.output_report
        Model_output_report_fig_index["svm"] = svm.output_report_fig_index


        for k in range(len(svm.output_report)):
            if k in svm.output_report_fig_index:
                output_report_doc.add_picture(svm.output_report[k])
            else:
                output_report_doc.add_paragraph(svm.output_report[k])

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's svm total end')

        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", svm.best_score)
        print("================================================================================")
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(svm.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
        
        
    if "knn" in Model:
        print("================================================================================")
        print(j,"Running KNN Model")
        output_report_doc.add_paragraph('======================================================================')
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running KNN Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin knn total")

        knn = knnModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, feature_selection = feature_selection, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["knn"] = knn.best_score
        BestEstimator["knn"] = knn.best_estimator
        Selected_X["knn"] = knn.selected_X
        Selected_columns_index["knn"] = knn.selected_columns_index
        Best_Params["knn"] = knn.best_params
        FI_features_names["knn"] = knn.features_names
        FI_imp["knn"] = knn.imp
        FI_name_index["knn"] = knn.name_index
        Model_output_report["knn"] = knn.output_report
        Model_output_report_fig_index["knn"] = knn.output_report_fig_index
        
        for k in range(len(knn.output_report)):
            if k in knn.output_report_fig_index:
                output_report_doc.add_picture(knn.output_report[k])
            else:
                output_report_doc.add_paragraph(knn.output_report[k])

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's knn total end')

        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", knn.best_score)
        print("================================================================================")
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(knn.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
        
    if "logReg" in Model:
        print("================================================================================")
        print(j,"Running Logistic Regression Model")
        output_report_doc.add_paragraph('======================================================================')
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running Logistic Regression Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin logReg total")

        logReg = logRegModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, feature_selection = feature_selection, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["logReg"] = logReg.best_score
        BestEstimator["logReg"] = logReg.best_estimator
        Selected_X["logReg"] = logReg.selected_X
        Selected_columns_index["logReg"] = logReg.selected_columns_index
        Best_Params["logReg"] = logReg.best_params
        FI_features_names["logReg"] = logReg.features_names
        FI_imp["logReg"] = logReg.imp
        FI_name_index["logReg"] = logReg.name_index
        Model_output_report["logReg"] = logReg.output_report
        Model_output_report_fig_index["logReg"] = logReg.output_report_fig_index
        
        for k in range(len(logReg.output_report)):
            if k in logReg.output_report_fig_index:
                output_report_doc.add_picture(logReg.output_report[k])
            else:
                output_report_doc.add_paragraph(logReg.output_report[k])

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's logReg total end')

        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", logReg.best_score)
        print("================================================================================")
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(logReg.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
        
    if "decisionTree" in Model:
        print("================================================================================")
        print(j,"Running Decision Tree Model")
        output_report_doc.add_paragraph('======================================================================')
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running Decision Tree Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin decisionTree total")

        decisionTree = decisionTreeModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, feature_selection = feature_selection, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["decisionTree"] = decisionTree.best_score
        BestEstimator["decisionTree"] = decisionTree.best_estimator
        Selected_X["decisionTree"] = decisionTree.selected_X
        Selected_columns_index["decisionTree"] = decisionTree.selected_columns_index
        Best_Params["decisionTree"] = decisionTree.best_params
        FI_features_names["decisionTree"] = decisionTree.features_names
        FI_imp["decisionTree"] = decisionTree.imp
        FI_name_index["decisionTree"] = decisionTree.name_index
        Model_output_report["decisionTree"] = decisionTree.output_report
        Model_output_report_fig_index["decisionTree"] = decisionTree.output_report_fig_index
        
        for k in range(len(decisionTree.output_report)):
            if k in decisionTree.output_report_fig_index:
                output_report_doc.add_picture(decisionTree.output_report[k])
            else:
                output_report_doc.add_paragraph(decisionTree.output_report[k])

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's decisionTree total end')


        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", decisionTree.best_score)
        print("================================================================================")
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(decisionTree.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
        tree_Graph = decisionTree.graphVisual()
        tree_Graph.format = 'png'
        tree_Graph.render('./pics/tree', view=False)
        output_report_doc.add_picture('./pics/tree.png')
#         display(decisionTree.graphVisual())
    if "randomForest" in Model:
        # close the feature_selection
        print("================================================================================")
        print(j,"Running Random Forest Model")
        output_report_doc.add_paragraph('======================================================================')
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running Random Forest Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin randomForest total")

        randomForest = randomForestModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, feature_selection = feature_selection, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["randomForest"] = randomForest.best_score
        BestEstimator["randomForest"] = randomForest.best_estimator
        Selected_X["randomForest"] = randomForest.selected_X
        Selected_columns_index["randomForest"] = randomForest.selected_columns_index
        Best_Params["randomForest"] = randomForest.best_params
        FI_features_names["randomForest"] = randomForest.features_names
        FI_imp["randomForest"] = randomForest.imp
        FI_name_index["randomForest"] = randomForest.name_index
        Model_output_report["randomForest"] = randomForest.output_report
        Model_output_report_fig_index["randomForest"] = randomForest.output_report_fig_index
        
        for k in range(len(randomForest.output_report)):
            if k in randomForest.output_report_fig_index:
                output_report_doc.add_picture(randomForest.output_report[k])
            else:
                output_report_doc.add_paragraph(randomForest.output_report[k])

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's randomForest total end')


        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", randomForest.best_score)
        print("================================================================================") 
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(randomForest.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
        
    if "bayes" in Model:
        print("================================================================================")
        print(j,"Running Gaussian Naive Bayes Model")
        output_report_doc.add_paragraph('======================================================================')
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running Gaussian Naive Bayes Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin bayes total")

        bayes = GaussianNBModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, feature_selection = feature_selection, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["bayes"] = bayes.best_score
        BestEstimator["bayes"] = bayes.best_estimator
        Selected_X["bayes"] = bayes.selected_X
        Selected_columns_index["bayes"] = bayes.selected_columns_index
        Best_Params["bayes"] = bayes.best_params
        FI_features_names["bayes"] = bayes.features_names
        FI_imp["bayes"] = bayes.imp
        FI_name_index["bayes"] = bayes.name_index
        Model_output_report["bayes"] = bayes.output_report
        Model_output_report_fig_index["bayes"] = bayes.output_report_fig_index
        
        for k in range(len(bayes.output_report)):
            if k in bayes.output_report_fig_index:
                output_report_doc.add_picture(bayes.output_report[k])
            else:
                output_report_doc.add_paragraph(bayes.output_report[k])

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's bayes total end')

        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", bayes.best_score)
        print("================================================================================")
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(bayes.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
        
    if "xgboost" in Model:
        print("================================================================================")
        print(j,"Running XGBoost Model")
        output_report_doc.add_paragraph('======================================================================')
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running XGBoost Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin xgboost total")

        xgboost = XGBoostModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, feature_selection = feature_selection, label_number = label_number, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["xgboost"] = xgboost.best_score
        BestEstimator["xgboost"] = xgboost.best_estimator
        Selected_X["xgboost"] = xgboost.selected_X
        Selected_columns_index["xgboost"] = xgboost.selected_columns_index
        Best_Params["xgboost"] = xgboost.best_params
        FI_features_names["xgboost"] = xgboost.features_names
        FI_imp["xgboost"] = xgboost.imp
        FI_name_index["xgboost"] = xgboost.name_index
        Model_output_report["xgboost"] = xgboost.output_report
        Model_output_report_fig_index["xgboost"] = xgboost.output_report_fig_index
        
        for k in range(len(xgboost.output_report)):
            if k in xgboost.output_report_fig_index:
                output_report_doc.add_picture(xgboost.output_report[k])
            else:
                output_report_doc.add_paragraph(xgboost.output_report[k])

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's xgboost total end')


        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", xgboost.best_score)
        print("================================================================================") 
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(xgboost.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
    if "stacking" in Model:
        #Feature imporatance of Stacking only for after modeling
        print("================================================================================")
        print(j,"Running Stacking Model")
        output_report_doc.add_paragraph('======================================================================')
        paragraph = output_report_doc.add_paragraph()
        paragraph.add_run('%d. Running Stacking Model' % (j)).bold = True
        j = j + 1

        time_start = time.time()
        random_num = random.randint(0, 99)
        output_report_doc.add_paragraph(str(random_num) + ". time record begin stacking total")

        stacking = StackingModel(scoringMetric=scoringMetric, iterations=iterations,X=X_scaled, y=y, 
                                 feature_selection = feature_selection, 
                                 BestEstimator = BestEstimator, Selected_columns_index = Selected_columns_index, evaluation = evaluation, curve_point=curve_point, pipeline_type=pipeline_type).runTest()
        BestScore["stacking"] = stacking.best_score
        BestEstimator["stacking"] = stacking.best_estimator
        Selected_X["stacking"] = stacking.selected_X
        Best_Params["stacking"] = stacking.best_params
        FI_features_names["stacking"] = stacking.features_names
        FI_imp["stacking"] = stacking.imp
        FI_name_index["stacking"] = stacking.name_index
        
        for k in range(len(stacking.output_report)):
            if k in stacking.output_report_fig_index:
                output_report_doc.add_picture(stacking.output_report[k])
            else:
                output_report_doc.add_paragraph(stacking.output_report[k])
        
#         Selected_columns_index["stacking"] = stacking.selected_columns_index

        time_end = time.time()
        output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's stacking total end')

        print("--------------------------------------------------------------------------------")
        print("Best ",scoringMetric," is ", stacking.best_score)
        print("================================================================================")
        output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
        tem_str = "Best "+ str(scoringMetric) + " is " + str(stacking.best_score)
        output_report_doc.add_paragraph(tem_str)
        output_report_doc.add_paragraph('======================================================================')
        
    print("Summary")
    print("================================================================================")
    print(scoringMetric," from Selected Models:")
    print("--------------------------------------------------------------------------------")
    paragraph = output_report_doc.add_paragraph()
    paragraph.add_run("Summary").bold = True
    output_report_doc.add_paragraph('======================================================================')
    tem_str = str(scoringMetric) + " from Selected Models:"
    output_report_doc.add_paragraph(tem_str)
    output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")


    final_no_list = []
    final_model_list = []
    final_score_list = []
    final_feature_no_list = []

    i = 0

    for index in BestScore:
        i = i + 1
        final_no_list.append(i)
        final_model_list.append(index)
        final_score_list.append(round(float(BestScore[index]), 5))
        final_feature_no_list.append(len(FI_imp[index]))

    COLUMNS = 4
    ROWS = len(BestScore)

    table = output_report_doc.add_table(rows=ROWS + 1, cols=COLUMNS, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'No.'
    hdr_cells[1].text = 'Model'
    hdr_cells[2].text = str(scoringMetric)
    hdr_cells[3].text = 'number of genes'
    list_0 = final_no_list
    list_1 = final_model_list
    list_2 = final_score_list
    list_3 = final_feature_no_list
    table_cells = table._cells

    for i in range(1, ROWS + 1):
        row_cells = table_cells[i * COLUMNS:(i + 1) * COLUMNS]
        # row_cells = table.add_row().cells
        row_cells[0].text = str(list_0[i - 1])
        row_cells[1].text = str(list_1[i - 1])
        row_cells[2].text = str(list_2[i - 1])
        row_cells[3].text = str(list_3[i - 1])




    print()

    if not final_evaluation:
        for key in Selected_X.keys():
                model_selected_X = Selected_X[key]
                model_selected_X.to_csv("./models/" + key + "_selected_X.csv", header=True, sep='\t')

        if test_file != False:
            print("================================================================================")
            print("Evaluate for Testset:")
            output_report_doc.add_paragraph('======================================================================')
            output_report_doc.add_paragraph("Evaluate for Testset:")
            for best_model in final_model_list:
                print("--------------------------------------------------------------------------------------------------------------------")
                print("Evaluate for " + best_model + ":")
                output_report_doc.add_paragraph('--------------------------------------------------------------------------------------------------------------------')
                output_report_doc.add_paragraph("Evaluate for " + best_model + ":")
                estimator = BestEstimator[best_model]
                evaluation_Model = clone(estimator)
                X_evaluation = Selected_X[best_model]
                print("X_evaluation shape :", X_evaluation)
                X_evaluation = X_evaluation.reindex(sorted(X_evaluation.columns), axis=1)
                evaluation_Model.fit(X_evaluation, y)
                cb_test = Circ_Biomarker(test_file).load_data()
                test_X = cb_test.X
                test_y = cb_test.y
                print(test_X.shape)

                selected_feature_list_input = []
                for feature_name in cb.X.columns:
                    selected_feature_list_input.append(str(feature_name))
                selected_test_X_input = test_X[test_X.columns.intersection(selected_feature_list_input)]

                test_X_scaled = selected_test_X_input

                selected_feature_list = []
                for feature_name in X_evaluation.columns:
                    selected_feature_list.append(str(feature_name))

                selected_test_X = test_X_scaled[test_X_scaled.columns.intersection(selected_feature_list)]
                selected_test_X = selected_test_X.reindex(sorted(selected_test_X.columns), axis=1)
                predcit_y = evaluation_Model.predict(selected_test_X)

                test_accuracy = accuracy_score(y_true=test_y, y_pred=predcit_y)
                test_balanced_accuracy = balanced_accuracy_score(y_true = test_y, y_pred = predcit_y)
                # test_average_precision = precision_score(y_true = test_y, y_pred = predcit_y, average='macro')
                test_f1 = f1_score(y_true = test_y, y_pred = predcit_y, average='macro')
                test_precision = precision_score(y_true = test_y, y_pred = predcit_y, average='macro')
                test_recall = recall_score(y_true = test_y, y_pred = predcit_y, average='macro')
                if label_number > 2:
                    y_one_hot = label_binarize(test_y, np.arange(label_number))
                    y_score = evaluation_Model.predict_proba(selected_test_X)
                    test_roc_auc = roc_auc_score(y_true = y_one_hot, y_score = y_score, average='macro', multi_class='ovo')
                else:
                    test_roc_auc = roc_auc_score(y_true = test_y, y_score = predcit_y, average='macro')

                print("Accuracy: %.3f%% " % (test_accuracy.mean() * 100.0))
                print("Balanced Accuracy: %.3f%% " % (test_balanced_accuracy.mean() * 100.0))
                # print("Precision_Recall_AUC: %.3f%% (%.3f%%)" % (
                # average_precision.mean() * 100.0, average_precision.std() * 100.0))
                print("ROC_AUC: %.3f%% " % (test_roc_auc.mean() * 100.0))
                print("Precision: %.3f%% " % (test_precision.mean() * 100.0))
                print("Recall / Sensitivity: %.3f%% " % (test_recall.mean() * 100.0))
                print("F1: %.3f%% " % (test_f1.mean() * 100.0))


                output_report_doc.add_paragraph("Accuracy: %.3f%% " % (test_accuracy.mean() * 100.0))
                output_report_doc.add_paragraph("Balanced Accuracy: %.3f%% " % (test_balanced_accuracy.mean() * 100.0))
                # output_report_doc.add_paragraph("Precision_Recall_AUC: %.3f%% (%.3f%%)" % (
                # average_precision.mean() * 100.0, average_precision.std() * 100.0))
                output_report_doc.add_paragraph("ROC_AUC: %.3f%% " % (test_roc_auc.mean() * 100.0))
                output_report_doc.add_paragraph("Precision: %.3f%% " % (test_precision.mean() * 100.0))
                output_report_doc.add_paragraph("Recall / Sensitivity: %.3f%% " % (test_recall.mean() * 100.0))
                output_report_doc.add_paragraph("F1: %.3f%% " % (test_f1.mean() * 100.0))

                best_model_feature_num = len(FI_imp[best_model])
                Result_df = pd.DataFrame(
                    {"BestModel": [best_model], "feature_num": [best_model_feature_num],
                     "Accuracy": [test_accuracy.mean()], #"Accuracy_std": [test_accuracy.std()],
                     "Balanced_Accuracy": [test_balanced_accuracy.mean()],#"Balanced_Accuracy_std": [test_balanced_accuracy.std()],
                     "ROC_AUC": [test_roc_auc.mean()], #"ROC_AUC_std": [test_roc_auc.std()],
                     "Precision": [test_precision.mean()], #"Precision_std": [test_precision.std()],
                     "Recall": [test_recall.mean()],#"Recall_std": [test_recall.std()],
                     "F1": [test_f1.mean()], #"F1_std": [test_f1.std()]
                     })
                Result_df.to_csv('./' + best_model + '_result.csv', sep=",", index=False)



                def feature_importance_final(best_model, output_report_doc, FI_features_names, FI_imp, FI_name_index, X, Model_cat="best_model"):
                    print(" ")

                    print("Feature Importance of the " + best_model)
                    print()
                    output_report_doc.add_paragraph(" ")

                    tem_str = "Feature Importance of the " + best_model
                    output_report_doc.add_paragraph(tem_str)

                    features_names = FI_features_names[best_model]
                    imp = FI_imp[best_model]
                    name_index = FI_name_index[best_model]
                    names = []

                    feature_name = []
                    feature_type = []
                    feature_score = []


                    for j in range(len(imp)):
                        names.append(features_names[name_index[j]])
                        print('Feature: %s, Score: %.5f' % (features_names[name_index[j]], imp[j]))

                        feature_name.append(features_names[name_index[j]])
                        feature_type.append("M")
                        feature_score.append(imp[j])
                        if best_model == "stacking" or best_model == "randomForest":
                            if float(imp[j]) == 0.0:
                                continue
                        output_report_doc.add_paragraph('Feature: %s, Score: %.5f' % (features_names[name_index[j]], imp[j]))

                    time_start = time.time()
                    random_num = random.randint(0, 99)
                    output_report_doc.add_paragraph(str(random_num) + ". time record begin correlation print")

                    threshold_show_last_step = 0.8

                    # output_report_doc.add_paragraph("======================================================================")
                    # tem_str = "High Correlation Features"
                    # output_report_doc.add_paragraph(tem_str)
                    # tem_str = "Threshold: " + str(threshold_show_last_step)
                    # output_report_doc.add_paragraph(tem_str)

                    correlation_matrix_final = X.corr()

                    threshold = threshold_show_last_step

                    for p in range(len(imp)):
                        for i in range(len(correlation_matrix_final.columns)):
                            if correlation_matrix_final.columns[i] == features_names[name_index[p]]:
                                flag = 0
                                for j in range(len(correlation_matrix_final.columns)):
                                    if (abs(correlation_matrix_final.iloc[i, j]) > threshold) and (flag == 0) and (i != j):
                                        # output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
                                        # output_report_doc.add_paragraph(str(features_names[name_index[p]]))
                                        # output_report_doc.add_paragraph('Feature: %s, Score: %.3f' % (
                                        #     correlation_matrix_final.columns[j], correlation_matrix_final.iloc[i, j]))

                                        feature_name.append(correlation_matrix_final.columns[j])
                                        feature_type.append("R")
                                        feature_score.append(correlation_matrix_final.iloc[i, j])

                                        flag += 1
                                    elif (abs(correlation_matrix_final.iloc[i, j]) > threshold) and (flag == 1) and (i != j):
                                        # output_report_doc.add_paragraph('Feature: %s, Score: %.3f' % (
                                        #     correlation_matrix_final.columns[j], correlation_matrix_final.iloc[i, j]))
                                        feature_name.append(correlation_matrix_final.columns[j])
                                        feature_type.append("R")
                                        feature_score.append(correlation_matrix_final.iloc[i, j])

                    time_end = time.time()
                    output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's correlation print end')

                    feature_imp_df = pd.DataFrame({"feature_name": feature_name,
                                              "feature_type": feature_type,
                                              "feature_score": feature_score})

                    feature_imp_df.to_csv("./" + best_model+"_imp_features.txt", sep="\t", index=False)





                    return output_report_doc

                output_report_doc = feature_importance_final(best_model, output_report_doc, FI_features_names, FI_imp,
                                                             FI_name_index, X, Model_cat="best_model")

            # if suggest_model_name_1 != best_model:
            #     output_report_doc = feature_importance_final(suggest_model_name_1, output_report_doc, FI_features_names, FI_imp,
            #                                                  FI_name_index, X, Model_cat="suggest_model_95")
            # if suggest_model_name_2 != best_model and suggest_model_name_2 != suggest_model_name_1:
            #     output_report_doc = feature_importance_final(suggest_model_name_2, output_report_doc, FI_features_names, FI_imp,
            #                                                  FI_name_index, X, Model_cat="suggest_model_90")



    print("================================================================================")
    output_report_doc.add_paragraph('======================================================================')
    output_report_doc.add_page_break()
    output_report_doc.save("./Report.doc")

    with open('./done.txt', 'w') as rsh:
        rsh.write("")


################Original Code######################

#     def select_interpretability_model(model_list):
#         model_interpretability_list = ["logReg", "decisionTree", "randomForest", "xgboost", "svm", "knn", "bayes",
#                                        "stacking"]
#         for model in model_interpretability_list:
#             if model in model_list:
#                 return model
#
#
#
#     best_model_score = max(final_score_list)
#
#     best_model_list = []
#     best_model_feature_no_list = []
#
#     for i in range(len(final_model_list)):
#         if final_score_list[i] == best_model_score:
#             best_model_list.append(final_model_list[i])
#             best_model_feature_no_list.append(final_feature_no_list[i])
#
#
#     if len(best_model_list) > 1:
#         best_model_feature_no_list_least = min(best_model_feature_no_list)
#         tem_list_best_model_select = []
#         for i in range(len(best_model_list)):
#             if best_model_feature_no_list[i] == best_model_feature_no_list_least:
#                 tem_list_best_model_select.append(best_model_list[i])
#         if len(tem_list_best_model_select) > 1:
#             best_model = select_interpretability_model(tem_list_best_model_select)
#         else:
#             best_model = tem_list_best_model_select[0]
#     else:
#         best_model = best_model_list[0]
#
#     best_model_index = final_model_list.index(best_model)
#     best_model_performance = final_score_list[best_model_index]
#     suggest_model_accept_threshold_1 = float(best_model_performance) * 0.95
#     suggest_model_accept_threshold_2 = float(best_model_performance) * 0.9
#
#     suggest_model_1_model_list = []
#     suggest_model_1_feature_no_list = []
#
#     suggest_model_2_model_list = []
#     suggest_model_2_feature_no_list = []
#
#     for i in range(len(final_model_list)):
#         if final_score_list[i] >= suggest_model_accept_threshold_1:
#             suggest_model_1_model_list.append(final_model_list[i])
#             suggest_model_1_feature_no_list.append(final_feature_no_list[i])
#         if final_score_list[i] >= suggest_model_accept_threshold_2:
#             suggest_model_2_model_list.append(final_model_list[i])
#             suggest_model_2_feature_no_list.append(final_feature_no_list[i])
#
#     def suggest_model(suggest_model_model_list, suggest_model_feature_no_list, best_model):
#         suggest_model_least_feature_no = min(suggest_model_feature_no_list)
#         suggest_model_model_list_selected = []
#         suggest_model_feature_no_list_selected = []
#         for i in range(len(suggest_model_feature_no_list)):
#             if suggest_model_feature_no_list[i] == suggest_model_least_feature_no:
#                 suggest_model_model_list_selected.append(suggest_model_model_list[i])
#                 suggest_model_feature_no_list_selected.append(suggest_model_feature_no_list[i])
#         if len(suggest_model_model_list_selected) > 1:
#             if best_model in suggest_model_model_list_selected:
#                 suggest_model_name = best_model
#             else:
#                 suggest_model_name = select_interpretability_model(suggest_model_model_list_selected)
#         else:
#             suggest_model_name = suggest_model_model_list_selected[0]
#
#         return suggest_model_name
#
#     suggest_model_name_1 = suggest_model(suggest_model_1_model_list, suggest_model_1_feature_no_list, best_model)
#     suggest_model_name_2 = suggest_model(suggest_model_2_model_list, suggest_model_2_feature_no_list, best_model)
#
#
#
#
#     for key in Selected_X.keys():
#         model_selected_X = Selected_X[key]
#         model_selected_X.to_csv("./models/" + key + "_selected_X.csv", header=True, sep='\t')
#
#
#     def final_evaluation(best_model,  output_report_doc, Best_Params, BestEstimator, Selected_X, evaluation_final, label_number, y, FI_imp, evaluation, Model_cat = "best_model"):
#         if best_model == "stacking":
#             print(Best_Params[best_model])
#             output_report_doc.add_paragraph(str(Best_Params[best_model]))
#         else:
#             print(BestEstimator[best_model])
#             output_report_doc.add_paragraph(str(BestEstimator[best_model]))
#
#         time_start = time.time()
#         random_num = random.randint(0, 99)
#         output_report_doc.add_paragraph(str(random_num) + ". time record begin evaluation")
#
#         print("--------------------------------------------------------------------------------")
#         print("Performance of the " + Model_cat + ": ", evaluation_final)
#         print()
#         output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
#         tem_str = "Performance of the " + Model_cat + ": " + ' '.join([str(elem) for elem in evaluation_final])
#         output_report_doc.add_paragraph(tem_str)
#     #     output_report_doc.add_paragraph(" ")
#
#
#
#         X_evaluation = Selected_X[best_model]
#
#         for method in evaluation_final:
#             if method == "cv":
#                 cvf = RepeatedStratifiedKFold(n_splits=10, n_repeats=100, random_state=12)
#             elif method == "headout":
#                 cvf = StratifiedShuffleSplit(n_splits=100, test_size=1/3, random_state=12)
#             elif method == "bootstrap":
#                 cvf = BootStrapSplit(n_splits=100, test_size=1/3, random_state=12)
#
#
#             estimator=BestEstimator[best_model]
#             from sklearn.metrics import make_scorer, precision_score, roc_auc_score
#             precision_scorer = make_scorer(precision_score, average='macro')
#             if label_number > 2 :
#                 roc_auc_scorer = 'roc_auc_ovo' #make_scorer(roc_auc_score, average='macro', multi_class = 'ovo')
#             else:
#                 roc_auc_scorer = 'roc_auc'
#
#             print("--------------------------------------------------------------------------------")
#             print("Evaluation Method: ", method)
#             output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
#             tem_str = "Evaluation Method: " + method
#             output_report_doc.add_paragraph(tem_str)
#
#             accuracy = cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring="accuracy", n_jobs=-1)
#             balanced_accuracy=cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring="balanced_accuracy", n_jobs=-1)
#         #     average_precision=cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring="average_precision_macro", n_jobs=-1)
#             average_precision = cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring=precision_scorer ,n_jobs=-1)
#             f1=cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring="f1_macro", n_jobs=-1)
#             precision=cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring="precision_macro", n_jobs=-1)
#             recall=cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring="recall_macro", n_jobs=-1)
#             # roc_auc=cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring="roc_auc", n_jobs=-1)
#             roc_auc = cross_val_score(estimator, X_evaluation, y, cv=cvf, scoring=roc_auc_scorer, n_jobs=-1)
#
#             print("Accuracy: %.3f%% (%.3f%%)" % (accuracy.mean() * 100.0, accuracy.std() * 100.0))
#             print("Balanced Accuracy: %.3f%% (%.3f%%)" % (balanced_accuracy.mean()*100.0, balanced_accuracy.std()*100.0))
#             print("Precision_Recall_AUC: %.3f%% (%.3f%%)" % (average_precision.mean()*100.0, average_precision.std()*100.0))
#             print("ROC_AUC: %.3f%% (%.3f%%)" % (roc_auc.mean()*100.0, roc_auc.std()*100.0))
#             print("Precision: %.3f%% (%.3f%%)" % (precision.mean()*100.0, precision.std()*100.0))
#             print("Recall / Sensitivity: %.3f%% (%.3f%%)" % (recall.mean()*100.0, recall.std()*100.0))
#             print("F1: %.3f%% (%.3f%%)" % (f1.mean()*100.0, f1.std()*100.0))
#             output_report_doc.add_paragraph("Accuracy: %.3f%% (%.3f%%)" % (accuracy.mean() * 100.0, accuracy.std() * 100.0))
#             output_report_doc.add_paragraph("Balanced Accuracy: %.3f%% (%.3f%%)" % (balanced_accuracy.mean()*100.0, balanced_accuracy.std()*100.0))
#             output_report_doc.add_paragraph("Precision_Recall_AUC: %.3f%% (%.3f%%)" % (average_precision.mean()*100.0, average_precision.std()*100.0))
#             output_report_doc.add_paragraph("ROC_AUC: %.3f%% (%.3f%%)" % (roc_auc.mean()*100.0, roc_auc.std()*100.0))
#             output_report_doc.add_paragraph("Precision: %.3f%% (%.3f%%)" % (precision.mean()*100.0, precision.std()*100.0))
#             output_report_doc.add_paragraph("Recall / Sensitivity: %.3f%% (%.3f%%)" % (recall.mean()*100.0, recall.std()*100.0))
#             output_report_doc.add_paragraph("F1: %.3f%% (%.3f%%)" % (f1.mean()*100.0, f1.std()*100.0))
#
#             if method == evaluation:
#                 best_model_feature_num = len(FI_imp[best_model])
#                 Result_df = pd.DataFrame({"Accuracy": [accuracy.mean()], "Balanced_Accuracy": [balanced_accuracy.mean()],
#                                       "ROC_AUC": [roc_auc.mean()], "Precision": [precision.mean()], "Recall": [recall.mean()],
#                                       "F1": [f1.mean()], "BestModel": [best_model], "feature_num": [best_model_feature_num]})
#                 Result_df.to_csv('./' + Model_cat + '_result.csv', sep=",", index = False)
#
#         time_end = time.time()
#         output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's evaluation')
#
#         return output_report_doc
#
#
#     print("--------------------------------------------------------------------------------")
#     print("The Best Model is ",best_model)
#     print("--------------------------------------------------------------------------------")
#     print("The Setting of the best model:")
#     print()
#     output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
#     tem_str = "The Best Model is " + best_model
#     output_report_doc.add_paragraph(tem_str)
#     output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
#     output_report_doc.add_paragraph("The Setting of the best model:")
# #     output_report_doc.add_paragraph(" ")
#
#     output_report_doc = final_evaluation(best_model, output_report_doc, Best_Params, BestEstimator, Selected_X,
#                                          evaluation_final, label_number, y, FI_imp, evaluation, Model_cat="best_model")
#     if suggest_model_name_1 != best_model:
#
#         print("--------------------------------------------------------------------------------")
#         print("The suggest_model less than 5% performance loss is ", suggest_model_name_1)
#         print("--------------------------------------------------------------------------------")
#         print("The Setting of the best model:")
#         print()
#         output_report_doc.add_paragraph(
#             "--------------------------------------------------------------------------------------------------------------------")
#         tem_str = "The suggest_model less than 5% performance loss is " + suggest_model_name_1
#         output_report_doc.add_paragraph(tem_str)
#         output_report_doc.add_paragraph(
#             "--------------------------------------------------------------------------------------------------------------------")
#         output_report_doc.add_paragraph("The Setting of the model:")
#
#         output_report_doc = final_evaluation(suggest_model_name_1, output_report_doc, Best_Params, BestEstimator, Selected_X,
#                                              evaluation_final, label_number, y, FI_imp, evaluation, Model_cat="suggest_model_95")
#     if suggest_model_name_2 != best_model and suggest_model_name_2 != suggest_model_name_1:
#
#         print("--------------------------------------------------------------------------------")
#         print("The suggest_model less than 10% performance loss is ", suggest_model_name_2)
#         print("--------------------------------------------------------------------------------")
#         print("The Setting of the best model:")
#         print()
#         output_report_doc.add_paragraph(
#             "--------------------------------------------------------------------------------------------------------------------")
#         tem_str = "The suggest_model less than 10% performance loss is " + suggest_model_name_2
#         output_report_doc.add_paragraph(tem_str)
#         output_report_doc.add_paragraph(
#             "--------------------------------------------------------------------------------------------------------------------")
#         output_report_doc.add_paragraph("The Setting of the model:")
#
#         output_report_doc = final_evaluation(suggest_model_name_2, output_report_doc, Best_Params, BestEstimator, Selected_X,
#                                              evaluation_final, label_number, y, FI_imp, evaluation, Model_cat="suggest_model_90")
#
#     if test_file != False:
#         print("================================================================================")
#         print("Evaluate for Testset:")
#         output_report_doc.add_paragraph('======================================================================')
#         output_report_doc.add_paragraph("Evaluate for Testset:")
#         estimator = BestEstimator[best_model]
#         evaluation_Model = clone(estimator)
#         X_evaluation = Selected_X[best_model]
#         print("X_evaluation shape :", X_evaluation)
#         X_evaluation = X_evaluation.reindex(sorted(X_evaluation.columns), axis=1)
#         evaluation_Model.fit(X_evaluation, y)
#         cb_test = Circ_Biomarker(test_file).load_data()
#         test_X = cb_test.X
#         test_y = cb_test.y
#         print(test_X.shape)
#
#         selected_feature_list_input = []
#         for feature_name in cb.X.columns:
#             selected_feature_list_input.append(str(feature_name))
#         selected_test_X_input = test_X[test_X.columns.intersection(selected_feature_list_input)]
#
#         test_X_scaled = selected_test_X_input
#
#         selected_feature_list = []
#         for feature_name in X_evaluation.columns:
#             selected_feature_list.append(str(feature_name))
#
#         selected_test_X = test_X_scaled[test_X_scaled.columns.intersection(selected_feature_list)]
#         selected_test_X = selected_test_X.reindex(sorted(selected_test_X.columns), axis=1)
#         predcit_y = evaluation_Model.predict(selected_test_X)
#
#         test_accuracy = accuracy_score(y_true=test_y, y_pred=predcit_y)
#         test_balanced_accuracy = balanced_accuracy_score(y_true = test_y, y_pred = predcit_y)
#         # test_average_precision = precision_score(y_true = test_y, y_pred = predcit_y, average='macro')
#         test_f1 = f1_score(y_true = test_y, y_pred = predcit_y, average='macro')
#         test_precision = precision_score(y_true = test_y, y_pred = predcit_y, average='macro')
#         test_recall = recall_score(y_true = test_y, y_pred = predcit_y, average='macro')
#         if label_number > 2:
#             y_one_hot = label_binarize(test_y, np.arange(label_number))
#             y_score = evaluation_Model.predict_proba(selected_test_X)
#             test_roc_auc = roc_auc_score(y_true = y_one_hot, y_score = y_score, average='macro', multi_class='ovo')
#         else:
#             test_roc_auc = roc_auc_score(y_true = test_y, y_score = predcit_y, average='macro')
#
#         print("Accuracy: %.3f%% (%.3f%%)" % (test_accuracy.mean() * 100.0, test_accuracy.std() * 100.0))
#         print("Balanced Accuracy: %.3f%% (%.3f%%)" % (test_balanced_accuracy.mean() * 100.0, test_balanced_accuracy.std() * 100.0))
#         # print("Precision_Recall_AUC: %.3f%% (%.3f%%)" % (
#         # average_precision.mean() * 100.0, average_precision.std() * 100.0))
#         print("ROC_AUC: %.3f%% (%.3f%%)" % (test_roc_auc.mean() * 100.0, test_roc_auc.std() * 100.0))
#         print("Precision: %.3f%% (%.3f%%)" % (test_precision.mean() * 100.0, test_precision.std() * 100.0))
#         print("Recall / Sensitivity: %.3f%% (%.3f%%)" % (test_recall.mean() * 100.0, test_recall.std() * 100.0))
#         print("F1: %.3f%% (%.3f%%)" % (test_f1.mean() * 100.0, test_f1.std() * 100.0))
#
#         output_report_doc.add_paragraph(
#             "Accuracy: %.3f%% (%.3f%%)" % (test_accuracy.mean() * 100.0, test_accuracy.std() * 100.0))
#         output_report_doc.add_paragraph(
#             "Balanced Accuracy: %.3f%% (%.3f%%)" % (test_balanced_accuracy.mean() * 100.0, test_balanced_accuracy.std() * 100.0))
#         # output_report_doc.add_paragraph("Precision_Recall_AUC: %.3f%% (%.3f%%)" % (
#         # average_precision.mean() * 100.0, average_precision.std() * 100.0))
#         output_report_doc.add_paragraph("ROC_AUC: %.3f%% (%.3f%%)" % (test_roc_auc.mean() * 100.0, test_roc_auc.std() * 100.0))
#         output_report_doc.add_paragraph(
#             "Precision: %.3f%% (%.3f%%)" % (test_precision.mean() * 100.0, test_precision.std() * 100.0))
#         output_report_doc.add_paragraph(
#             "Recall / Sensitivity: %.3f%% (%.3f%%)" % (test_recall.mean() * 100.0, test_recall.std() * 100.0))
#         output_report_doc.add_paragraph("F1: %.3f%% (%.3f%%)" % (test_f1.mean() * 100.0, test_f1.std() * 100.0))
#
#
#     def feature_importance_final(best_model, output_report_doc, FI_features_names, FI_imp, FI_name_index, X, Model_cat="best_model"):
#         print("======================================================================")
#         print("Feature Importance of the " + Model_cat)
#         print()
#         output_report_doc.add_paragraph("======================================================================")
#         tem_str = "Feature Importance of the " + Model_cat
#         output_report_doc.add_paragraph(tem_str)
#
#         features_names = FI_features_names[best_model]
#         imp = FI_imp[best_model]
#         name_index = FI_name_index[best_model]
#         names = []
#
#         feature_name = []
#         feature_type = []
#         feature_score = []
#
#
#         for j in range(len(imp)):
#             names.append(features_names[name_index[j]])
#             print('Feature: %s, Score: %.5f' % (features_names[name_index[j]], imp[j]))
#             feature_name.append(features_names[name_index[j]])
#             feature_type.append("M")
#             feature_score.append(imp[j])
#             output_report_doc.add_paragraph('Feature: %s, Score: %.5f' % (features_names[name_index[j]], imp[j]))
#
#         time_start = time.time()
#         random_num = random.randint(0, 99)
#         output_report_doc.add_paragraph(str(random_num) + ". time record begin correlation print")
#
#         threshold_show_last_step = 0.8
#
#         output_report_doc.add_paragraph("======================================================================")
#         tem_str = "High Correlation Features"
#         output_report_doc.add_paragraph(tem_str)
#         tem_str = "Threshold: " + str(threshold_show_last_step)
#         output_report_doc.add_paragraph(tem_str)
#
#         correlation_matrix_final = X.corr()
#
#         threshold = threshold_show_last_step
#
#         for p in range(len(imp)):
#             for i in range(len(correlation_matrix_final.columns)):
#                 if correlation_matrix_final.columns[i] == features_names[name_index[p]]:
#                     flag = 0
#                     for j in range(len(correlation_matrix_final.columns)):
#                         if (abs(correlation_matrix_final.iloc[i, j]) > threshold) and (flag == 0) and (i != j):
#                             output_report_doc.add_paragraph("--------------------------------------------------------------------------------------------------------------------")
#                             output_report_doc.add_paragraph(str(features_names[name_index[p]]))
#                             output_report_doc.add_paragraph('Feature: %s, Score: %.3f' % (
#                                 correlation_matrix_final.columns[j], correlation_matrix_final.iloc[i, j]))
#
#                             feature_name.append(correlation_matrix_final.columns[j])
#                             feature_type.append("R")
#                             feature_score.append(correlation_matrix_final.iloc[i, j])
#
#                             flag += 1
#                         elif (abs(correlation_matrix_final.iloc[i, j]) > threshold) and (flag == 1) and (i != j):
#                             output_report_doc.add_paragraph('Feature: %s, Score: %.3f' % (
#                                 correlation_matrix_final.columns[j], correlation_matrix_final.iloc[i, j]))
#                             feature_name.append(correlation_matrix_final.columns[j])
#                             feature_type.append("R")
#                             feature_score.append(correlation_matrix_final.iloc[i, j])
#
#         time_end = time.time()
#         output_report_doc.add_paragraph(str(random_num) + '. time cost: ' + str(time_end - time_start) + 's correlation print end')
#
#         feature_imp_df = pd.DataFrame({"feature_name": feature_name,
#                                   "feature_type": feature_type,
#                                   "feature_score": feature_score})
#
#         feature_imp_df.to_csv("./" + Model_cat+"_imp_features.txt", sep="\t", index=False)
#         return output_report_doc
#
#     output_report_doc = feature_importance_final(best_model, output_report_doc, FI_features_names, FI_imp,
#                                                  FI_name_index, X, Model_cat="best_model")
#
#     if suggest_model_name_1 != best_model:
#         output_report_doc = feature_importance_final(suggest_model_name_1, output_report_doc, FI_features_names, FI_imp,
#                                                      FI_name_index, X, Model_cat="suggest_model_95")
#     if suggest_model_name_2 != best_model and suggest_model_name_2 != suggest_model_name_1:
#         output_report_doc = feature_importance_final(suggest_model_name_2, output_report_doc, FI_features_names, FI_imp,
#                                                      FI_name_index, X, Model_cat="suggest_model_90")
#
#
#
#     print("================================================================================")
#     output_report_doc.add_paragraph('======================================================================')
#     output_report_doc.add_page_break()
#     output_report_doc.save("./Report.doc")

    ################Original Code End######################

    
    return BestEstimator, Selected_X, y, Selected_columns_index

def main():
    parser = argparse.ArgumentParser(formatter_class=argparse.ArgumentDefaultsHelpFormatter)
    parser.add_argument("--input", default=None, type=str, required=True, help="Path of the input data.")
    parser.add_argument("--info_table", default=None, type=str, required=True, help="Path of the info table.")
    parser.add_argument("--model", nargs='+', default=["svm"],
                        help="Model Selection" 
                            "choices=[svm, knn, logReg, decisionTree, randomForest, bayes, xgboost, stacking]")
    parser.add_argument("--metric", choices=["F1", "AUC","ACC", "BACC"], default="F1",
                        help="Scoring Metric" )
    parser.add_argument("--val_method_selection", choices=["cv", "headout", "bootstrap"], default="cv", help="Evaluation Method in model selection" )
    # parser.add_argument("--val_method_final", choices=["cv", "headout", "bootstrap"], default="cv", help="Evaluation Method in Final score ")
    parser.add_argument("--val_method_final", nargs='+', default=["cv"],
                        help="Evaluation Method in Final score "
                        "choices=[cv, headout, bootstrap]")

    parser.add_argument("--feature_selection", type=str, default="True", help="Scoring Metric" )

    parser.add_argument("--testset", type=str, default="False", help="whether use test set evaluation" )

    parser.add_argument("--corrthreshold", type=float, default=0.8, help="Threshold for high correlation remove")
    parser.add_argument("--curve_point", choices=["knee", "highest"], default="knee", help="the metric to select the point in feature number curve" )

    parser.add_argument("--pipeline_type", choices=["general", "fresh"], default="general", help="type of pipeline")

    parser.add_argument("--linc_padj", type=str, default="False",
                        help="lincRNA DESeq2 output file path(include padj information)")
    parser.add_argument("--PC_padj", type=str, default="False",
                        help="Protein Coding DESeq2 output file path(include padj information)")
    parser.add_argument("--linc_FPKM", type=str, default="False",
                        help="lincRNA FPKM file path")
    parser.add_argument("--PC_FPKM", type=str, default="False",
                        help="Protein Coding FPKM file path")
    parser.add_argument("--Circ_padj", type=str, default="False",
                        help="CircRNA file path(include padj information)")
    parser.add_argument("--Circ_FPKM", type=str, default="False",
                        help="CircRNA FPKM file path")

    parser.add_argument("--microarray_padj", type=str, default="False",
                        help="microarray limma output file path(include padj information)")
    parser.add_argument("--microarray_FPKM", type=str, default="False",
                        help="microarray FPKM file path")


    parser.add_argument("--pseudo_foldchange", type=float, default=float(1e-6), help="Pseudo-count for abs log2 foldchange calculation")
    parser.add_argument("--pvalue_type", choices=["padj", "pvalue"], default="padj", help="p-value type for different expressions")
    parser.add_argument("--pvalue_threshold", type=float, default=0.05, help="p-value threshold for different expressions")
    parser.add_argument("--foldchange_threshold", type=float, default=1.0, help="abs log2 foldchange threshold for different expressions")

    parser.add_argument("--PC_maxmin_remove_threshold", type=float, default=1.0, help="remove threshold for minimum of Protein maximum")
    parser.add_argument("--linc_maxmin_remove_threshold", type=float, default=0.001, help="remove threshold for minimum of lincRNA maximum")
    parser.add_argument("--Circ_maxmin_remove_threshold", type=float, default=0, help="remove threshold for minimum of Circular RNA maximum")
    parser.add_argument("--microarray_maxmin_remove_threshold", type=float, default=0, help="remove threshold for minimum of microarray maximum")


    parser.add_argument("--final_evaluation", type=str, default="False", help="whether it is the final evaluation")



    args = parser.parse_args()

    def t_or_f(fs):
        ua = str(fs).upper()
        if 'TRUE'.startswith(ua):
            return True
        elif 'FALSE'.startswith(ua):
            return False
        else:
            return True

    def t_or_f_testset(fs):
        ua = str(fs).upper()
        if 'FALSE'.startswith(ua):
            return False
        else:
            return fs
    feature_selection_requirment = t_or_f(args.feature_selection)
    testset_file = t_or_f_testset(args.testset)

    linc_padj_path = t_or_f_testset(args.linc_padj)
    PC_padj_path = t_or_f_testset(args.PC_padj)
    linc_FPKM_path = t_or_f_testset(args.linc_FPKM)
    PC_FPKM_path = t_or_f_testset(args.PC_FPKM)
    Circ_padj_path = t_or_f_testset(args.Circ_padj)
    Circ_FPKM_path = t_or_f_testset(args.Circ_FPKM)
    microarray_padj_path = t_or_f_testset(args.microarray_padj)
    microarray_FPKM_path = t_or_f_testset(args.microarray_FPKM)

    final_evaluation = t_or_f(args.final_evaluation)


    #feature_selection BestEstimator, Selected_X, y, Selected_columns_index = circBiomarker(input_file, scoringMetric="F1", feature_selection=True, Model=["knn"])

    BestEstimator, Selected_X, y, Selected_columns_index = circBiomarker(input_file = args.input, info_table_path = args.info_table, scoringMetric=args.metric,
                                                                         feature_selection=feature_selection_requirment, Model=args.model,
                                                                         test_file=testset_file, val_method_selection=args.val_method_selection,
                                                                         val_method_final = args.val_method_final, corrthreshold=args.corrthreshold, curve_point = args.curve_point, pipeline_type = args.pipeline_type,
                                                                         linc_padj=linc_padj_path, PC_padj=PC_padj_path, linc_FPKM=linc_FPKM_path, PC_FPKM=PC_FPKM_path,
                                                                         Circ_padj=Circ_padj_path, Circ_FPKM=Circ_FPKM_path,pseudo_foldchange=args.pseudo_foldchange,
                                                                         pvalue_type=args.pvalue_type, pvalue_threshold = args.pvalue_threshold, foldchange_threshold = args.foldchange_threshold,
                                                                         PC_maxmin_remove_threshold = args.PC_maxmin_remove_threshold, linc_maxmin_remove_threshold = args.linc_maxmin_remove_threshold,
                                                                         Circ_maxmin_remove_threshold = args.Circ_maxmin_remove_threshold, final_evaluation=final_evaluation,
                                                                         microarray_padj=microarray_padj_path, microarray_FPKM=microarray_FPKM_path,
                                                                         microarray_maxmin_remove_threshold=args.microarray_maxmin_remove_threshold)

if __name__ == '__main__':
    if not sys.warnoptions:
        warnings.simplefilter("ignore")
        os.environ["PYTHONWARNINGS"] = "ignore"  # Also affect subprocesses
    main()